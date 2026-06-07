"""
tab3_design.py — Tab 3: Design
Rigid Pavement Design V7
AASHTO 1993 — คำนวณความหนาแผ่นคอนกรีต JPCP/JRCP & CRCP
"""
import streamlit as st
import streamlit.components.v1 as components
import matplotlib.pyplot as plt
from io import BytesIO
from datetime import datetime

from engine import (
    calc_w18, check_design, find_optimum_k, compare_d,
    convert_cube_to_cyl, calc_ec, calc_sc, get_zr,
    plot_structure, fig_to_bytes,
    D_PAIRS,
)

# ── สีหลัก ───────────────────────────────────────────────────
_JPCP_BD   = '#1565C0'
_JPCP_BDLT = '#90CAF9'
_JPCP_BG   = '#E3F2FD'
_CRCP_BD   = '#2E7D32'
_CRCP_BDLT = '#A5D6A7'
_CRCP_BG   = '#E8F5E9'

SC_FIXED = 600.0   # psi — กรมทางหลวง กำหนด max


# ============================================================
# UI Helpers
# ============================================================
def _card_header(text, color):
    st.markdown(
        f'<div style="background:{color};border-radius:6px 6px 0 0;'
        f'padding:6px 12px;font-size:12px;font-weight:700;color:#fff;'
        f'margin-bottom:0">{text}</div>',
        unsafe_allow_html=True)

def _row(label, value, hi=False):
    c = _JPCP_BD if hi else '#1A237E'
    st.markdown(
        f'<div style="display:flex;justify-content:space-between;'
        f'padding:3px 0;border-bottom:1px solid rgba(0,0,0,0.06);font-size:12px">'
        f'<span style="color:#78909C">{label}</span>'
        f'<span style="font-family:IBM Plex Mono,monospace;font-weight:600;color:{c}">'
        f'{value}</span></div>', unsafe_allow_html=True)

def _mbox(label, value, unit='', vc='#1565C0', bg='#E3F2FD'):
    st.markdown(
        f'<div style="background:{bg};border:1px solid rgba(0,0,0,0.08);'
        f'border-radius:7px;padding:8px;text-align:center;margin-bottom:4px">'
        f'<div style="font-size:10px;color:#78909C;margin-bottom:2px">{label}</div>'
        f'<div style="font-family:IBM Plex Mono,monospace;font-size:20px;'
        f'font-weight:700;color:{vc}">{value}</div>'
        f'<div style="font-size:10px;color:#78909C">{unit}</div></div>',
        unsafe_allow_html=True)

def _verdict_bar(d_cm, d_in, w18_cap, w18_req, ratio, passed, bd_color):
    pct_cap   = min(ratio * 100, 100)
    bar_color = '#43A047' if passed else '#E53935'
    label     = f'✅ ผ่าน  (×{ratio:.2f})' if passed else f'❌ ไม่ผ่าน  (×{ratio:.2f})'
    ratio_txt = f'+{(ratio-1)*100:.0f}%' if passed else f'{ratio*100:.0f}%'
    st.markdown(
        f'<div style="background:#F5F5F5;border:1px solid {bd_color}33;'
        f'border-radius:8px;padding:8px 10px;margin-bottom:4px">'
        f'<div style="display:flex;justify-content:space-between;font-size:12px;margin-bottom:4px">'
        f'<span style="font-family:IBM Plex Mono,monospace;font-weight:700;color:{bd_color}">'
        f'D = {d_in} in ({d_cm} ซม.)</span>'
        f'<span style="font-weight:700;color:{bar_color}">{label}</span></div>'
        f'<div style="position:relative;background:#E0E0E0;border-radius:4px;height:10px">'
        f'<div style="background:{bar_color};width:{pct_cap:.1f}%;height:10px;border-radius:4px;'
        f'opacity:0.85"></div>'
        f'<div style="position:absolute;top:0;left:0;width:100%;height:100%;'
        f'border-right:2px dashed #9E9E9E;border-radius:4px;pointer-events:none"></div>'
        f'</div>'
        f'<div style="display:flex;justify-content:space-between;font-size:10px;'
        f'color:#90A4AE;margin-top:3px">'
        f'<span>W18_cap = {w18_cap:,.0f}</span>'
        f'<span style="color:{bar_color};font-weight:600">{ratio_txt} จาก W18_req</span>'
        f'<span>W18_req = {w18_req:,.0f}</span>'
        f'</div></div>',
        unsafe_allow_html=True)

def _kopt_box(prefix, rec_d_cm, k_opt, k_eff, bd):
    if k_opt is None:
        return
    delta  = k_eff - k_opt
    ok     = k_eff >= k_opt
    bg     = _CRCP_BG if ok else '#FFEBEE'
    bc     = _CRCP_BDLT if ok else '#EF9A9A'
    vc     = _CRCP_BD if ok else '#C62828'
    symbol = '✅' if ok else '⚠️'
    margin = f'{delta:+.0f} pci ({delta/k_opt*100:+.1f}%)'
    st.markdown(
        f'<div style="background:{bg};border:2px solid {bc};border-radius:8px;'
        f'padding:10px 12px;margin-top:6px">'
        f'<div style="font-size:12px;font-weight:700;color:{vc};margin-bottom:6px">'
        f'{symbol} k_opt vs k_eff  —  D = {rec_d_cm} ซม. ({round(rec_d_cm/2.54)} in)</div>'
        f'<div style="display:flex;gap:8px">'
        f'<div style="flex:1;background:white;border-radius:6px;padding:6px;text-align:center">'
        f'<div style="font-size:10px;color:#78909C">k_opt (min required)</div>'
        f'<div style="font-family:IBM Plex Mono,monospace;font-size:18px;font-weight:700;color:{bd}">'
        f'{k_opt:.0f} pci</div></div>'
        f'<div style="flex:1;background:white;border-radius:6px;padding:6px;text-align:center">'
        f'<div style="font-size:10px;color:#78909C">k_eff (Tab 2)</div>'
        f'<div style="font-family:IBM Plex Mono,monospace;font-size:18px;font-weight:700;color:{vc}">'
        f'{k_eff:.0f} pci</div></div>'
        f'<div style="flex:1;background:white;border-radius:6px;padding:6px;text-align:center">'
        f'<div style="font-size:10px;color:#78909C">Δk = k_eff − k_opt</div>'
        f'<div style="font-family:IBM Plex Mono,monospace;font-size:14px;font-weight:700;color:{vc}">'
        f'{margin}</div></div>'
        f'</div></div>',
        unsafe_allow_html=True)


# ============================================================
# Word Report (ไม่เปลี่ยน logic)
# ============================================================
def _create_word_report(ptype, proj_name, params, rows, sel_d_cm,
                        struct_fig_bytes, date_str,
                        fig33_bytes=None, fig34_bytes=None):
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return None

    TH = 'TH SarabunPSK'
    EQ = 'Times New Roman'
    TS = Pt(15)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = TH
    style.font.size = TS

    sec = doc.sections[0]
    sec.page_width  = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin  = sec.bottom_margin = Cm(2.5)

    h0 = doc.add_heading('รายการคำนวณออกแบบความหนาถนนคอนกรีต', 0)
    h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('ตามวิธี AASHTO 1993')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = TH; p.runs[0].font.size = TS

    doc.add_heading('1. ข้อมูลทั่วไป', level=1)
    w18_src = '(กรอกเอง)' if params.get('w18_manual') else '(จาก ESAL JSON)'
    for txt in [
        f'ชื่อโครงการ: {proj_name or "—"}',
        f'ประเภทถนน: {ptype}',
        f'วันที่คำนวณ: {date_str}',
        f'แหล่งข้อมูล W18: {w18_src}',
    ]:
        p = doc.add_paragraph(txt)
        p.runs[0].font.name = TH; p.runs[0].font.size = TS

    doc.add_heading('2. ข้อมูลนำเข้า', level=1)
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    for i, h in enumerate(['พารามิเตอร์', 'สัญลักษณ์', 'ค่า', 'หน่วย']):
        c = t.rows[0].cells[i]
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.name = TH; r.font.size = TS

    input_rows = [
        ('ESAL ออกแบบ',              'W₁₈',    f"{params['w18']:,.0f}",    'ESALs'),
        ('Terminal Serviceability',  'Pt',      f"{params['pt']:.1f}",      '—'),
        ('Reliability',              'R',       f"{params['R']:.0f}",       '%'),
        ('Standard Deviation',       'So',      f"{params['so']:.2f}",      '—'),
        ('k_eff (Tab 2)',             'k_eff',   f"{params['k_eff']:,.0f}", 'pci'),
        ("กำลังอัด (Cube)",           "f'c",     f"{params['fc_cube']:.0f}",'ksc'),
        ("กำลังอัด (Cylinder)",       "f'c,cyl", f"{params['fc_cyl']:.0f}",'ksc'),
        ('Modulus of Rupture',        'Sc',      f"{params['sc']:.0f}",     'psi'),
        ('Modulus of Elasticity',     'Ec',      f"{params['ec']:,.0f}",    'psi'),
        ('Load Transfer Coefficient', 'J',       f"{params['j']:.1f}",      '—'),
        ('Drainage Coefficient',      'Cd',      f"{params['cd']:.1f}",     '—'),
        ('ΔPSI',                      'ΔPSI',    f"{params['dpsi']:.1f}",   '—'),
    ]
    for param, sym, val, unit in input_rows:
        row = t.add_row().cells
        for i, txt in enumerate([param, sym, val, unit]):
            r = row[i].paragraphs[0].add_run(txt)
            r.font.name = TH; r.font.size = TS

    doc.add_heading('3. สมการออกแบบ AASHTO 1993', level=1)

    def _add_eq(document, parts):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '360')
        pPr.append(ind)
        for text, is_sub, is_sup in parts:
            run = p.add_run(text)
            run.font.name = EQ; run.font.size = Pt(13)
            if is_sub or is_sup:
                rPr = run._r.get_or_add_rPr()
                va = OxmlElement('w:vertAlign')
                va.set(qn('w:val'), 'subscript' if is_sub else 'superscript')
                rPr.append(va)
        return p

    _add_eq(doc, [
        ('log', False, False), ('10', True, False),
        ('(W', False, False),  ('18', True, False),
        (') = Z', False, False), ('R', False, False),
        (' \u00d7 S', False, False), ('o', True, False),
        (' + 7.35 \u00d7 log', False, False), ('10', True, False),
        ('(D+1) \u2212 0.06', False, False),
    ])
    _add_eq(doc, [
        ('        + log', False, False), ('10', True, False),
        ('(\u0394PSI/3.0) / (1 + 1.624\u00d710', False, False),
        ('7', False, True), ('/(D+1)', False, False),
        ('8.46', False, True), (')', False, False),
    ])
    _add_eq(doc, [
        ('        + (4.22 \u2212 0.32\u00d7P', False, False),
        ('t', True, False), (') \u00d7 log', False, False),
        ('10', True, False), ('[(S', False, False),
        ('c', True, False), ('\u00d7C', False, False),
        ('d', True, False), ('\u00d7(D', False, False),
        ('0.75', False, True), ('\u22121.132))/(215.63\u00d7J\u00d7(D', False, False),
        ('0.75', False, True), ('\u221218.42/(E', False, False),
        ('c', True, False), ('/k)', False, False),
        ('0.25', False, True), (')]', False, False),
    ])

    p = doc.add_paragraph('โดยที่:')
    p.runs[0].font.name = TH; p.runs[0].font.size = TS
    tsym = doc.add_table(rows=1, cols=3)
    tsym.style = 'Table Grid'
    for i, h in enumerate(['สัญลักษณ์', 'ความหมาย', 'หน่วย']):
        c = tsym.rows[0].cells[i]
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.name = TH; r.font.size = TS
    for sym, meaning, unit in [
        ('W18',  'จำนวน ESAL ที่รองรับได้',                       'ESALs'),
        ('ZR',   'Standard Normal Deviate',                        '-'),
        ('So',   'Overall Standard Deviation',                     '-'),
        ('D',    'ความหนาแผ่นคอนกรีต',                             'in'),
        ('DPSI', 'การสูญเสีย Serviceability (4.5 - Pt)',           '-'),
        ('Sc',   'Modulus of Rupture',                             'psi'),
        ('Cd',   'Drainage Coefficient',                           '-'),
        ('J',    'Load Transfer Coefficient',                      '-'),
        ('Ec',   'Modulus of Elasticity of Concrete',              'psi'),
        ('k',    'Effective Modulus of Subgrade Reaction (k_eff)', 'pci'),
    ]:
        row = tsym.add_row().cells
        for i, txt in enumerate([sym, meaning, unit]):
            r = row[i].paragraphs[0].add_run(txt)
            r.font.name = TH; r.font.size = TS

    doc.add_paragraph()
    doc.add_heading('4. ผลการเปรียบเทียบความหนา', level=1)
    t3 = doc.add_table(rows=1, cols=7)
    t3.style = 'Table Grid'
    for i, h in enumerate(['D (ซม.)', 'D (นิ้ว)', 'W18 ต้องการ',
                            'log10(W18_cap)', 'W18 รองรับได้', 'อัตราส่วน', 'ผล']):
        c = t3.rows[0].cells[i]
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.name = TH; r.font.size = TS
    for rv in rows:
        row = t3.add_row().cells
        for i, txt in enumerate([
            f"{rv['d_cm']}",
            f"{rv['d_inch']}",
            f"{rv.get('w18_req', params['w18']):,.0f}",
            f"{rv['log_w18']:.4f}",
            f"{rv['w18_cap']:,.0f}",
            f"{rv['ratio']:.3f}",
            'ผ่าน' if rv['passed'] else 'ไม่ผ่าน',
        ]):
            r = row[i].paragraphs[0].add_run(txt)
            r.font.name = TH; r.font.size = TS

    k_opt = params.get('k_opt')
    k_eff = params.get('k_eff')
    if k_opt and k_eff:
        doc.add_heading('5. k_opt vs k_eff', level=1)
        delta   = k_eff - k_opt
        verdict = 'เพียงพอ' if delta >= 0 else 'ไม่เพียงพอ'
        tk = doc.add_table(rows=1, cols=2)
        tk.style = 'Table Grid'
        for i, h in enumerate(['รายการ', 'ค่า']):
            c = tk.rows[0].cells[i]
            r = c.paragraphs[0].add_run(h)
            r.bold = True; r.font.name = TH; r.font.size = TS
        for lbl, val in [
            (f"D แนะนำ",                f"{sel_d_cm} ซม. ({round(sel_d_cm/2.54)} in)"),
            ("k_opt (minimum required)", f"{k_opt:.0f} pci"),
            ("k_eff (จาก Tab 2)",        f"{k_eff:.0f} pci"),
            (f"Δk = k_eff - k_opt",     f"{delta:+.0f} pci  →  {verdict}"),
        ]:
            row = tk.add_row().cells
            for i, txt in enumerate([lbl, val]):
                r = row[i].paragraphs[0].add_run(txt)
                r.font.name = TH; r.font.size = TS

    doc.add_heading('6. สรุปผล', level=1)
    sel_row = next((r for r in rows if r['d_cm'] == sel_d_cm), None)
    summary = [f"ความหนาที่เลือก: {sel_d_cm} ซม. ({round(sel_d_cm/2.54)} นิ้ว)",
               f"ESAL ที่ต้องการ: {params['w18']:,.0f} ESALs"]
    if sel_row:
        w18_req_sel = sel_row.get('w18_req', params['w18'])
        summary += [
            f"ESAL ที่ต้องการ (D={sel_d_cm} ซม.): {w18_req_sel:,.0f} ESALs",
            f"ESAL ที่รองรับได้: {sel_row['w18_cap']:,.0f} ESALs",
            f"อัตราส่วน (capacity/demand): {sel_row['ratio']:.3f}",
            f"ผลการตรวจสอบ: {'ผ่านเกณฑ์' if sel_row['passed'] else 'ไม่ผ่านเกณฑ์'}",
        ]
    for txt in summary:
        p = doc.add_paragraph(txt)
        p.runs[0].font.name = TH; p.runs[0].font.size = TS

    if fig33_bytes:
        doc.add_heading('7. AASHTO Figure 3.3 — Composite k_inf', level=1)
        p = doc.add_paragraph('ผลการหาค่า Composite Modulus of Subgrade Reaction (k_inf) จาก Nomograph:')
        p.runs[0].font.name = TH; p.runs[0].font.size = TS
        doc.add_picture(BytesIO(fig33_bytes), width=Inches(5.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if fig34_bytes:
        doc.add_heading('8. AASHTO Figure 3.4 — Loss of Support', level=1)
        p = doc.add_paragraph('ผลการปรับแก้ k_eff ตาม Loss of Support (LS):')
        p.runs[0].font.name = TH; p.runs[0].font.size = TS
        doc.add_picture(BytesIO(fig34_bytes), width=Inches(5.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if struct_fig_bytes:
        doc.add_heading('9. รูปตัดโครงสร้างชั้นทาง', level=1)
        p = doc.add_paragraph('รูปตัดโครงสร้างชั้นทางที่ออกแบบ:')
        p.runs[0].font.name = TH; p.runs[0].font.size = TS
        doc.add_picture(BytesIO(struct_fig_bytes), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    ref = doc.add_paragraph('Reference: AASHTO Guide for Design of Pavement Structures 1993')
    ref.runs[0].font.name = TH; ref.runs[0].font.size = Pt(13)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ============================================================
# Design Block
# ============================================================
def _design_block(prefix, ptype, fc_cyl, ec_psi, cd, w18_req, pt, zr, so, bd, bdlt):
    dpsi  = 4.5 - pt
    k_eff = st.session_state.get(f'{prefix}_k_eff')

    if w18_req is None:
        st.markdown(
            '<div style="background:#FFF3E0;border:1px solid #FFB74D;border-radius:8px;'
            'padding:8px 12px;font-size:12px;color:#E65100">'
            '⚠️ ยังไม่มีข้อมูล W18 — กรุณากรอก W18 ในช่องด้านบน'
            '</div>', unsafe_allow_html=True)
        return None

    if k_eff is None:
        st.markdown(
            '<div style="background:#FFF3E0;border:1px solid #FFB74D;border-radius:8px;'
            'padding:8px 12px;font-size:12px;color:#E65100">'
            f'⚠️ ยังไม่มีค่า k_eff ({ptype}) — กรุณาคำนวณใน Tab 2 ก่อน'
            '</div>', unsafe_allow_html=True)
        return None

    if prefix == 'jpcp':
        j_opts  = [2.5, 2.6, 2.7, 2.8]
        j_def   = st.session_state.get('jpcp_j', 2.8)
        j_label = 'J — Load Transfer Coefficient (JPCP/JRCP)'
    else:
        j_opts  = [2.3, 2.4, 2.5, 2.6]
        j_def   = st.session_state.get('crcp_j', 2.6)
        j_label = 'J — Load Transfer Coefficient (CRCP)'

    if j_def not in j_opts:
        j_def = j_opts[-1]

    j_val = st.select_slider(
        j_label, options=j_opts,
        value=j_def,
        key=f'{prefix}_j',
        format_func=lambda x: f'{x:.1f}')

    ed = st.session_state.get('esal_data')
    w18_manual_mode = st.session_state.get('w18_manual_mode', False)

    w18_per_d = {}
    if ed and not w18_manual_mode:
        from engine import compute_esal_for_d
        for d_in, d_cm in D_PAIRS:
            w18_d, _, _ = compute_esal_for_d(
                ed['traffic_data'], pt,
                ed['lane_factor'], ed['direction_factor'], d_cm)
            w18_per_d[d_cm] = w18_d
    else:
        for d_in, d_cm in D_PAIRS:
            w18_per_d[d_cm] = w18_req
        st.markdown(
            '<div style="background:#FFF8E1;border:1px solid #FFD54F;'
            'border-radius:7px;padding:6px 10px;font-size:11px;color:#E65100;margin-bottom:4px">'
            '⚠️ W18 กรอกเอง — ใช้ค่าเดียวกันทุก D (conservative)'
            '</div>', unsafe_allow_html=True)

    st.markdown('<div style="height:4px"></div>', unsafe_allow_html=True)
    w18_note = 'แยกตาม D' if (ed and not w18_manual_mode) else 'ค่าเดียว (manual)'
    _row(f'W18 ({w18_note})', f'{w18_req:,.0f} ESALs (ref D=30)')
    _row('k_eff (Tab 2)',  f'{k_eff:.0f} pci')
    _row("f'c (cube)", f"{st.session_state.get('fc_cube', 350):.0f} ksc")
    _row('Ec',             f'{ec_psi:,.0f} psi')
    _row('Sc (ทล. lock)', f'{SC_FIXED:.0f} psi')
    _row('J',              f'{j_val:.1f}', hi=True)
    _row('Cd',             f'{cd:.1f}',    hi=True)
    _row('Pt / ΔPSI',      f'{pt:.1f} / {dpsi:.1f}')
    _row('ZR / So',        f'{zr:.3f} / {so:.2f}')
    st.markdown('<div style="height:6px"></div>', unsafe_allow_html=True)

    from engine import calc_w18 as _calc_w18
    rows = []
    for d_in, d_cm in D_PAIRS:
        w18_d  = w18_per_d[d_cm]
        lw, wc = _calc_w18(d_in, dpsi, pt, zr, so,
                           SC_FIXED, cd, j_val, ec_psi, k_eff)
        passed = wc >= w18_d
        ratio  = round(wc / w18_d, 3) if w18_d > 0 else 0
        rows.append({
            'd_cm':    d_cm,
            'd_inch':  d_in,
            'log_w18': round(lw, 4),
            'w18_cap': round(wc, 0),
            'w18_req': w18_d,
            'passed':  passed,
            'ratio':   ratio,
        })
    st.session_state[f'{prefix}_design_rows'] = rows

    passed_rows = [r for r in rows if r['passed']]
    for r in rows:
        _verdict_bar(r['d_cm'], r['d_inch'],
                     r['w18_cap'], r['w18_req'],
                     r['ratio'], r['passed'], bd)

    if passed_rows:
        rec = min(passed_rows, key=lambda r: r['d_cm'])
        _mbox(f'✅ D แนะนำ ({ptype})',
              f"{rec['d_inch']} in ({rec['d_cm']} ซม.)",
              f"W18 capacity = {rec['w18_cap']:,.0f}",
              _CRCP_BD if prefix == 'crcp' else _JPCP_BD,
              _CRCP_BG if prefix == 'crcp' else _JPCP_BG)
        st.session_state[f'{prefix}_rec_d_cm'] = rec['d_cm']
    else:
        st.markdown(
            '<div style="background:#FFEBEE;border:1px solid #EF9A9A;'
            'border-radius:8px;padding:8px 12px;font-size:12px;color:#C62828">'
            '❌ ไม่มี D ที่ผ่านเกณฑ์ในช่วง 25–35 ซม. — พิจารณาเพิ่ม k_eff หรือลด J'
            '</div>', unsafe_allow_html=True)
        st.session_state[f'{prefix}_rec_d_cm'] = None

    sel_d_cm = st.session_state.get(f'{prefix}_rec_d_cm') or 30
    sel_d_in = round(sel_d_cm / 2.54)
    k_opt    = find_optimum_k(w18_req, sel_d_in, dpsi, pt, zr, so,
                               SC_FIXED, cd, j_val, ec_psi)
    _kopt_box(prefix, sel_d_cm, k_opt, k_eff, bd)

    st.session_state[f'{prefix}_design_params'] = {
        'w18':        w18_req,
        'w18_manual': st.session_state.get('w18_manual_mode', False),
        'pt':         pt,
        'R':          st.session_state.get('reliability', 90),
        'so':         so,
        'k_eff':      k_eff,
        'fc_cube':    st.session_state.get('fc_cube', 350),
        'fc_cyl':     fc_cyl,
        'sc':         SC_FIXED,
        'ec':         ec_psi,
        'j':          j_val,
        'cd':         cd,
        'dpsi':       dpsi,
        'k_opt':      k_opt,
    }

    return {'rows': rows, 'j': j_val, 'k_eff': k_eff, 'k_opt': k_opt}


# ============================================================
# MAIN RENDER
# ============================================================
def render_tab3():
    ed = st.session_state.get('esal_data')
    pt = st.session_state.get('pt',  2.0)
    so = st.session_state.get('so',  0.35)
    R  = st.session_state.get('reliability', 90)
    zr = get_zr(R)

    w18_from_json = None
    if ed:
        from engine import compute_esal_for_d
        w18_from_json, _, _ = compute_esal_for_d(
            ed['traffic_data'], pt,
            ed['lane_factor'], ed['direction_factor'], 30)

    kj = st.session_state.get('jpcp_k_eff')
    kc = st.session_state.get('crcp_k_eff')

    # ════════════════════════════════════════════════════════
    # Card 1 — Status bar
    # ════════════════════════════════════════════════════════
    with st.container(border=True):
        st.markdown('<div class="rp-card-title">📋 สถานะข้อมูลจาก Tab 1-2</div>',
                    unsafe_allow_html=True)
        s1, s2, s3 = st.columns(3)
        with s1:
            if ed:
                st.markdown(
                    f'<div class="rp-status-ok">✅ W18 = {w18_from_json:,.0f} (จาก JSON)</div>',
                    unsafe_allow_html=True)
            else:
                st.markdown(
                    '<div class="rp-status-warn">⚠️ ไม่มี JSON — กรอก W18 เองด้านล่าง</div>',
                    unsafe_allow_html=True)
        with s2:
            if kj:
                st.markdown(
                    f'<div class="rp-status-ok">✅ k_eff JPCP = {kj:.0f} pci</div>',
                    unsafe_allow_html=True)
            else:
                st.markdown(
                    '<div class="rp-status-warn">⚠️ ยังไม่มี k_eff JPCP (Tab 2)</div>',
                    unsafe_allow_html=True)
        with s3:
            if kc:
                st.markdown(
                    f'<div class="rp-status-ok">✅ k_eff CRCP = {kc:.0f} pci</div>',
                    unsafe_allow_html=True)
            else:
                st.markdown(
                    '<div class="rp-status-warn">⚠️ ยังไม่มี k_eff CRCP (Tab 2)</div>',
                    unsafe_allow_html=True)

    # ════════════════════════════════════════════════════════
    # Card 2 — W18 Input
    # ════════════════════════════════════════════════════════
    with st.container(border=True):
        st.markdown('<div class="rp-card-title">🔢 W18 — ESAL ออกแบบ</div>',
                    unsafe_allow_html=True)
        if ed:
            use_manual = st.checkbox(
                'กรอก W18 เองแทน (override จาก JSON)',
                value=st.session_state.get('w18_manual_mode', False),
                key='w18_manual_mode')
            if use_manual:
                w18_req = st.number_input(
                    'W18 (ESALs)', min_value=100_000, max_value=500_000_000,
                    value=st.session_state.get('w18_manual', int(w18_from_json)),
                    step=100_000, key='w18_manual', format='%d')
                st.caption(f'W18 จาก JSON = {w18_from_json:,.0f} ESALs (ไม่ได้ใช้)')
            else:
                w18_req = w18_from_json
                st.markdown(
                    f'<div class="rp-status-ok">'
                    f'✅ ใช้ W18 = <b>{w18_req:,.0f} ESALs</b> จาก ESAL JSON (D = 30 ซม.)'
                    f'</div>', unsafe_allow_html=True)
        else:
            st.session_state['w18_manual_mode'] = True
            w18_req = st.number_input(
                'W18 (ESALs) — กรอกเอง', min_value=100_000, max_value=500_000_000,
                value=st.session_state.get('w18_manual', 5_000_000),
                step=100_000, key='w18_manual', format='%d',
                help='กรอก ESAL ออกแบบโดยตรง (ไม่มี JSON)')

    # ════════════════════════════════════════════════════════
    # Card 3 — Shared Parameters
    # ════════════════════════════════════════════════════════
    with st.container(border=True):
        st.markdown('<div class="rp-card-title">⚙️ พารามิเตอร์ร่วม (JPCP & CRCP)</div>',
                    unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        with c1:
            fc_cube = st.number_input(
                "f'c กำลังคอนกรีต Cube (ksc)", 280, 600,
                st.session_state.get('fc_cube', 350), step=10,
                key='fc_cube',
                help='กรมทางหลวง กำหนดไม่ต่ำกว่า 350 ksc')
            if fc_cube < 350:
                st.warning('⚠️ ต่ำกว่า 350 ksc — ไม่เป็นไปตามมาตรฐาน ทล.')
        fc_cyl = convert_cube_to_cyl(fc_cube)
        ec_psi = calc_ec(fc_cyl)
        with c2:
            st.markdown(
                f'<div style="background:{_JPCP_BG};border:1px solid {_JPCP_BDLT};'
                f'border-radius:8px;padding:8px;text-align:center;margin-top:4px">'
                f'<div style="font-size:10px;color:#78909C">f\'c,cyl = 0.8 × f\'c,cube = {fc_cyl:.0f} ksc</div>'
                f'<div style="font-family:IBM Plex Mono,monospace;font-size:18px;'
                f'font-weight:700;color:{_JPCP_BD}">{ec_psi:,.0f} psi</div>'
                f'<div style="font-size:10px;color:#78909C">Modulus of Elasticity (Ec)</div></div>',
                unsafe_allow_html=True)
        with c3:
            st.markdown(
                f'<div style="background:{_CRCP_BG};border:1px solid {_CRCP_BDLT};'
                f'border-radius:8px;padding:8px;text-align:center;margin-top:4px">'
                f'<div style="font-size:10px;color:#78909C">Sc — ทล. กำหนด (lock)</div>'
                f'<div style="font-family:IBM Plex Mono,monospace;font-size:18px;'
                f'font-weight:700;color:{_CRCP_BD}">{SC_FIXED:.0f} psi</div>'
                f'<div style="font-size:10px;color:#78909C">Modulus of Rupture</div></div>',
                unsafe_allow_html=True)
        st.markdown('<div style="height:6px"></div>', unsafe_allow_html=True)
        st.markdown('<div style="font-size:12px;color:#546E7A;margin-bottom:2px">'
                    'Cd — Drainage Coefficient (ใช้ร่วมกัน)</div>',
                    unsafe_allow_html=True)
        cd_str = st.radio(
            'Cd',
            options=['1.0 — ระบายน้ำปกติ', '1.1 — ระบายน้ำดี', '1.2 — ระบายน้ำดีมาก'],
            index=[1.0, 1.1, 1.2].index(st.session_state.get('cd', 1.0)),
            key='cd_radio',
            horizontal=True,
            label_visibility='collapsed')
        cd = float(cd_str.split(' — ')[0])
        st.session_state['cd'] = cd

    # ════════════════════════════════════════════════════════
    # Side-by-side Design (JPCP | CRCP)
    # ════════════════════════════════════════════════════════
    col_j, col_c = st.columns(2)

    with col_j:
        _card_header('🔲  JPCP / JRCP — Design', _JPCP_BD)
        with st.container(border=True):
            res_j = _design_block('jpcp', 'JPCP/JRCP',
                                  fc_cyl, ec_psi, cd,
                                  w18_req, pt, zr, so,
                                  _JPCP_BD, _JPCP_BDLT)
            if res_j and st.session_state.get('jpcp_layers'):
                if st.button('🏗️ โครงสร้าง JPCP', key='str_j', use_container_width=True):
                    st.session_state['jpcp_show_str3'] = \
                        not st.session_state.get('jpcp_show_str3', False)
                if st.session_state.get('jpcp_show_str3'):
                    rec_cm = st.session_state.get('jpcp_rec_d_cm')
                    fig = plot_structure(
                        st.session_state['jpcp_layers'],
                        concrete_cm=rec_cm,
                        title=f'JPCP  D = {rec_cm} cm' if rec_cm else 'JPCP Structure')
                    if fig:
                        st.pyplot(fig, use_container_width=True)
                        st.session_state['jpcp_struct_bytes'] = fig_to_bytes(fig)
                        st.download_button('⬇️ PNG โครงสร้าง JPCP',
                                           st.session_state['jpcp_struct_bytes'],
                                           'struct_jpcp.png', 'image/png',
                                           key='dl_str_j')
                        plt.close(fig)

    with col_c:
        _card_header('〰️  CRCP — Design', _CRCP_BD)
        with st.container(border=True):
            res_c = _design_block('crcp', 'CRCP',
                                  fc_cyl, ec_psi, cd,
                                  w18_req, pt, zr, so,
                                  _CRCP_BD, _CRCP_BDLT)
            if res_c and st.session_state.get('crcp_layers'):
                if st.button('🏗️ โครงสร้าง CRCP', key='str_c', use_container_width=True):
                    st.session_state['crcp_show_str3'] = \
                        not st.session_state.get('crcp_show_str3', False)
                if st.session_state.get('crcp_show_str3'):
                    rec_cm = st.session_state.get('crcp_rec_d_cm')
                    fig = plot_structure(
                        st.session_state['crcp_layers'],
                        concrete_cm=rec_cm,
                        title=f'CRCP  D = {rec_cm} cm' if rec_cm else 'CRCP Structure')
                    if fig:
                        st.pyplot(fig, use_container_width=True)
                        st.session_state['crcp_struct_bytes'] = fig_to_bytes(fig)
                        st.download_button('⬇️ PNG โครงสร้าง CRCP',
                                           st.session_state['crcp_struct_bytes'],
                                           'struct_crcp.png', 'image/png',
                                           key='dl_str_c')
                        plt.close(fig)

    # ════════════════════════════════════════════════════════
    # Comparison Table + PDF Export
    # ════════════════════════════════════════════════════════
    st.markdown('---')
    _comparison_table(res_j, res_c, fc_cube, ec_psi, cd, pt, zr, so)


# ============================================================
# PDF Summary Export (fpdf2)
# ============================================================
def _create_pdf_summary(proj_name, date_str, sections, layers_j, layers_c, dj_cm, dc_cm):
    """สร้าง PDF summary — ชื่อโครงการ + comparison table + layer table"""
    try:
        from fpdf import FPDF
    except ImportError:
        return None

    import os
    BASE_DIR  = os.path.dirname(os.path.abspath(__file__))
    FONT_REG  = os.path.join(BASE_DIR, 'Sarabun-Regular.ttf')
    FONT_BOLD = os.path.join(BASE_DIR, 'Sarabun-Bold.ttf')
    if not os.path.exists(FONT_REG):
        FONT_REG  = 'Sarabun-Regular.ttf'
        FONT_BOLD = 'Sarabun-Bold.ttf'

    # ── ดึง designer name (default '—') ──────────────────────
    designer = '—'

    class PDF(FPDF):
        def header(self):
            pass  # ไม่ใช้ auto-header — วาด header เองด้านล่าง

        def footer(self):
            self.set_y(-10)
            self.set_font('Sarabun', '', 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 8, f'Page {self.page_no()} | KMUTNB - ภาควิชาครุศาสตร์โยธา - มจพ.',
                      align='C')
            self.set_text_color(0, 0, 0)

    pdf = PDF(orientation='P', unit='mm', format='A4')
    pdf.add_font('Sarabun', '',  FONT_REG,  uni=True)
    pdf.add_font('Sarabun', 'B', FONT_BOLD, uni=True)
    pdf.set_font('Sarabun', '', 10)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=12)

    # ══════════════════════════════════════════════════════════
    # Header — Flexible style, สีน้ำเงิน
    # ══════════════════════════════════════════════════════════
    BLUE      = (21, 101, 192)   # #1565C0
    PAGE_W    = 190              # usable width (A4 210 - margin 10*2)
    H_TOP     = 10               # y เริ่มต้น header
    COL_R_W   = 55               # ความกว้างคอลัมน์ขวา (ผู้ออกแบบ/วันที่)
    COL_L_W   = PAGE_W - COL_R_W

    # ── ซ้าย: ชื่อรายงาน ──────────────────────────────────────
    pdf.set_xy(10, H_TOP)
    pdf.set_font('Sarabun', 'B', 14)
    pdf.set_text_color(*BLUE)
    pdf.cell(COL_L_W, 7, 'Rigid Pavement Design Report', ln=False)

    # ── ขวา: ผู้ออกแบบ ────────────────────────────────────────
    pdf.set_xy(10 + COL_L_W, H_TOP)
    pdf.set_font('Sarabun', '', 8)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(COL_R_W, 7, f'ผู้ออกแบบ: {designer}', align='R', ln=True)

    # ── ซ้าย: subtitle ────────────────────────────────────────
    pdf.set_xy(10, H_TOP + 7)
    pdf.set_font('Sarabun', '', 8)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(COL_L_W, 5, 'AASHTO 1993 · ภาควิชาครุศาสตร์โยธา มจพ.', ln=False)

    # ── ขวา: วันที่ ───────────────────────────────────────────
    pdf.set_xy(10 + COL_L_W, H_TOP + 7)
    pdf.set_font('Sarabun', '', 8)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(COL_R_W, 5, f'วันที่: {date_str}', align='R', ln=True)

    # ── เส้นคั่นสีน้ำเงิน ────────────────────────────────────
    pdf.set_draw_color(*BLUE)
    pdf.set_line_width(0.6)
    pdf.line(10, H_TOP + 13, 200, H_TOP + 13)
    pdf.set_line_width(0.2)
    pdf.set_draw_color(0, 0, 0)

    # ── Project name ──────────────────────────────────────────
    pdf.set_xy(10, H_TOP + 16)
    pdf.set_font('Sarabun', 'B', 10)
    pdf.set_text_color(*BLUE)
    pdf.cell(0, 6, f'Project: {proj_name or "(ไม่ระบุชื่อโครงการ)"}', ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(2)

    # ── helper: section header ────────────────────────────────
    def sec_header(title):
        pdf.set_fill_color(238, 242, 247)
        pdf.set_font('Sarabun', 'B', 8)
        pdf.set_text_color(84, 110, 122)
        pdf.cell(0, 5, f'  {title}', ln=True, fill=True)
        pdf.set_text_color(0, 0, 0)

    # ── helper: table row ────────────────────────────────────
    W_LABEL = 62
    W_COL   = 64

    def tbl_header():
        pdf.set_font('Sarabun', 'B', 9)
        pdf.set_fill_color(21, 101, 192)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(W_LABEL, 6, 'รายการ', border=0, fill=True)
        pdf.set_fill_color(21, 101, 192)
        pdf.cell(W_COL, 6, '  JPCP / JRCP', border=0, fill=True)
        pdf.set_fill_color(46, 125, 50)
        pdf.cell(W_COL, 6, '  CRCP', border=0, fill=True, ln=True)
        pdf.set_text_color(0, 0, 0)

    def tbl_row(label, val_j, val_c, shade=False, bold_val=False):
        pdf.set_fill_color(250, 250, 250) if shade else pdf.set_fill_color(255, 255, 255)
        pdf.set_font('Sarabun', '', 8)
        pdf.set_text_color(84, 110, 122)
        pdf.cell(W_LABEL, 5, f'  {label}', border='B', fill=True)
        pdf.set_text_color(26, 35, 126)
        f = 'B' if bold_val else ''
        pdf.set_font('Sarabun', f, 8)
        pdf.cell(W_COL, 5, f'  {val_j}', border='B', fill=True)
        pdf.set_text_color(27, 94, 32)
        pdf.cell(W_COL, 5, f'  {val_c}', border='B', fill=True, ln=True)
        pdf.set_text_color(0, 0, 0)

    def tbl_row_shared(label, val):
        """shared value — แสดงซ้ำทั้ง 2 column สีเทา"""
        pdf.set_fill_color(250, 250, 250)
        pdf.set_font('Sarabun', '', 8)
        pdf.set_text_color(84, 110, 122)
        pdf.cell(W_LABEL, 5, f'  {label}', border='B', fill=True)
        pdf.set_text_color(120, 120, 120)
        pdf.cell(W_COL, 5, f'  {val}', border='B', fill=True)
        pdf.cell(W_COL, 5, f'  {val}', border='B', fill=True, ln=True)
        pdf.set_text_color(0, 0, 0)

    # ── Comparison Table ─────────────────────────────────────
    tbl_header()

    for sec_title, rows in sections:
        sec_header(sec_title)
        for row in rows:
            if row.get('shared'):
                tbl_row_shared(row['label'], row['val_j'])
            else:
                tbl_row(row['label'], str(row['val_j']), str(row['val_c']),
                        shade=row.get('shade', False),
                        bold_val=row.get('bold', False))

    pdf.ln(4)

    # ── Layer Structure Table ─────────────────────────────────
    if layers_j or layers_c:
        pdf.set_font('Sarabun', 'B', 9)
        pdf.set_fill_color(21, 101, 192)
        pdf.set_text_color(255, 255, 255)
        W_NO  = 10
        W_MAT = 115
        W_LC  = 32
        pdf.cell(W_NO,  6, '#',               border=0, fill=True)
        pdf.cell(W_MAT, 6, '  วัสดุ',         border=0, fill=True)
        pdf.set_fill_color(21, 101, 192)
        pdf.cell(W_LC,  6, '  JPCP (ซม.)',    border=0, fill=True)
        pdf.set_fill_color(46, 125, 50)
        pdf.cell(W_LC,  6, '  CRCP (ซม.)',    border=0, fill=True, ln=True)
        pdf.set_text_color(0, 0, 0)

        # แผ่นคอนกรีต
        pdf.set_font('Sarabun', 'B', 9)
        pdf.set_fill_color(238, 242, 247)
        pdf.set_text_color(21, 101, 192)
        pdf.cell(W_NO,  6, '0',                      border='B', fill=True)
        pdf.cell(W_MAT, 6, '  แผ่นคอนกรีต (D)',     border='B', fill=True)
        pdf.cell(W_LC,  6, f'  {dj_cm or "-"}',      border='B', fill=True)
        pdf.set_text_color(46, 125, 50)
        pdf.cell(W_LC,  6, f'  {dc_cm or "-"}',      border='B', fill=True, ln=True)
        pdf.set_text_color(0, 0, 0)

        names_j = [l['name'] for l in layers_j]
        names_c = [l['name'] for l in layers_c]
        all_names = list(dict.fromkeys(names_j + names_c))

        def _thick(layers, name):
            for l in layers:
                if l['name'] == name:
                    return l['thickness_cm']
            return None

        tot_j = dj_cm or 0
        tot_c = dc_cm or 0
        for i, name in enumerate(all_names, 1):
            tj = _thick(layers_j, name)
            tc = _thick(layers_c, name)
            if tj: tot_j += tj
            if tc: tot_c += tc
            pdf.set_font('Sarabun', '', 8)
            pdf.set_fill_color(255, 255, 255)
            pdf.set_text_color(100, 100, 100)
            pdf.cell(W_NO, 6, str(i), border='B', fill=True)
            pdf.set_text_color(84, 110, 122)
            pdf.cell(W_MAT, 6, f'  {name[:65]}', border='B', fill=True)
            pdf.set_text_color(26, 35, 126)
            pdf.cell(W_LC, 6, f'  {tj if tj else "-"}', border='B', fill=True)
            pdf.set_text_color(27, 94, 32)
            pdf.cell(W_LC, 6, f'  {tc if tc else "-"}', border='B', fill=True, ln=True)
            pdf.set_text_color(0, 0, 0)

        # total row
        pdf.set_font('Sarabun', 'B', 9)
        pdf.set_fill_color(240, 244, 255)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(W_NO,  6, '',                           border='B', fill=True)
        pdf.cell(W_MAT, 6, '  รวมทั้งหมด (รวมคอนกรีต)', border='B', fill=True)
        pdf.set_text_color(21, 101, 192)
        pdf.cell(W_LC,  6, f'  {tot_j} ซม.',             border='B', fill=True)
        pdf.set_text_color(46, 125, 50)
        pdf.cell(W_LC,  6, f'  {tot_c} ซม.',             border='B', fill=True, ln=True)

    from io import BytesIO as _BytesIO
    buf = _BytesIO(pdf.output())
    buf.seek(0)
    return buf


# ============================================================
# Comparison Table — เรียกจาก render_tab3
# ============================================================
def _comparison_table(res_j, res_c, fc_cube, ec_psi, cd, pt, zr, so):
    """แสดง comparison table JPCP vs CRCP"""

    pj = st.session_state.get('jpcp_design_params', {})
    pc = st.session_state.get('crcp_design_params', {})
    kj_eff  = st.session_state.get('jpcp_k_eff')
    kc_eff  = st.session_state.get('crcp_k_eff')
    kj_opt  = pj.get('k_opt')
    kc_opt  = pc.get('k_opt')
    # ── เพิ่ม CBR และ MR_psi ─────────────────────────────────
    cbr     = st.session_state.get('cbr', 4.0)
    MR_psi  = st.session_state.get('MR_psi', 6000)
    # ─────────────────────────────────────────────────────────
    dj_cm   = st.session_state.get('jpcp_rec_d_cm')
    dc_cm   = st.session_state.get('crcp_rec_d_cm')
    rows_j  = st.session_state.get('jpcp_design_rows', [])
    rows_c  = st.session_state.get('crcp_design_rows', [])
    layers_j = st.session_state.get('jpcp_layers', [])
    layers_c = st.session_state.get('crcp_layers', [])

    if not (res_j or res_c):
        return

    def _get_row(rows, d_cm):
        return next((r for r in rows if r['d_cm'] == d_cm), None)

    rj  = _get_row(rows_j, dj_cm) if dj_cm else None
    rc_ = _get_row(rows_c, dc_cm) if dc_cm else None

    # ── CSS ──────────────────────────────────────────────────
    TABLE_CSS = """
    <style>
    .cmp-table { width:100%; border-collapse:collapse; font-size:12px; }
    .cmp-table th { padding:7px 10px; font-size:11px; font-weight:700; }
    .cmp-table td { padding:5px 10px; border-bottom:0.5px solid #F0F0F0; vertical-align:middle; }
    .cmp-th-label { text-align:left; color:#78909C; background:#FAFAFA; width:36%; }
    .cmp-th-j { background:#1565C0; color:#fff; text-align:center; width:32%; }
    .cmp-th-c { background:#2E7D32; color:#fff; text-align:center; width:32%; }
    .cmp-td-label { color:#546E7A; background:#FAFAFA; font-size:11px; }
    .cmp-td-b { text-align:center; font-family:'IBM Plex Mono',monospace; font-size:12px; color:#1A237E; }
    .cmp-td-g { text-align:center; font-family:'IBM Plex Mono',monospace; font-size:12px; color:#1B5E20; }
    .cmp-td-shared { text-align:center; color:#546E7A; font-family:'IBM Plex Mono',monospace; font-size:12px; }
    .cmp-shdr td { background:#EEF2F7; font-size:10px; font-weight:700; color:#546E7A;
                   text-transform:uppercase; letter-spacing:0.07em; padding:5px 10px; }
    .cmp-shared-tag { background:#F3E5F5; color:#6A1B9A; border-radius:3px;
                      font-size:9px; padding:1px 5px; margin-left:4px; }
    .cmp-badge-ok   { background:#E8F5E9; color:#2E7D32; border-radius:4px;
                      padding:3px 10px; font-size:11px; font-weight:700; display:inline-block; }
    .cmp-badge-fail { background:#FFEBEE; color:#C62828; border-radius:4px;
                      padding:3px 10px; font-size:11px; font-weight:700; display:inline-block; }
    .cmp-bar-wrap { background:#E0E0E0; border-radius:3px; height:5px; margin-top:4px; }
    .cmp-bar-b { background:#1565C0; border-radius:3px; height:5px; }
    .cmp-bar-g { background:#2E7D32; border-radius:3px; height:5px; }
    .cmp-layer-table { width:100%; border-collapse:collapse; font-size:11px; }
    .cmp-layer-table th { padding:5px 8px; font-size:10px; font-weight:700; }
    .cmp-layer-table td { padding:4px 8px; border-bottom:0.5px solid #F5F5F5; vertical-align:middle; }
    .cmp-lth-no  { background:#FAFAFA; color:#90A4AE; width:5%;  text-align:center; }
    .cmp-lth-mat { background:#FAFAFA; color:#78909C; width:42%; }
    .cmp-lth-j   { background:#E3F2FD; color:#1565C0; text-align:center; width:26.5%; font-size:10px; font-weight:700; }
    .cmp-lth-c   { background:#E8F5E9; color:#2E7D32; text-align:center; width:26.5%; font-size:10px; font-weight:700; }
    .cmp-ltd-no  { text-align:center; color:#B0BEC5; font-size:10px; }
    .cmp-ltd-mat { color:#546E7A; }
    .cmp-ltd-j   { text-align:center; font-family:'IBM Plex Mono',monospace; color:#1565C0; font-weight:600; }
    .cmp-ltd-c   { text-align:center; font-family:'IBM Plex Mono',monospace; color:#2E7D32; font-weight:600; }
    .cmp-ltd-na  { text-align:center; color:#BDBDBD; font-size:10px; }
    .cmp-total   { background:#F0F4FF; font-weight:700; }
    </style>
    """
    st.markdown(TABLE_CSS, unsafe_allow_html=True)

    def _badge(ok):
        if ok is None: return '<span style="color:#90A4AE">—</span>'
        cls = 'cmp-badge-ok' if ok else 'cmp-badge-fail'
        txt = '✅ ผ่าน' if ok else '❌ ไม่ผ่าน'
        return f'<span class="{cls}">{txt}</span>'

    def _val(v, fmt='{:.0f}', fallback='—'):
        return fmt.format(v) if v is not None else fallback

    def _bar(pct, color_class):
        w = min(pct * 100, 100) if pct else 0
        return (f'<div class="cmp-bar-wrap">'
                f'<div class="{color_class}" style="width:{w:.0f}%"></div></div>')

    dpsi  = 4.5 - pt
    dj_in = round(dj_cm / 2.54) if dj_cm else '—'
    dc_in = round(dc_cm / 2.54) if dc_cm else '—'

    # ════════════════════════════════════════════════════════
    # Card 1 — Design Comparison
    # ════════════════════════════════════════════════════════
    with st.container(border=True):
        st.markdown('<div class="rp-card-title">📊 เปรียบเทียบผลการออกแบบ — JPCP/JRCP vs CRCP</div>',
                    unsafe_allow_html=True)

        shared = '<span class="cmp-shared-tag">ร่วมกัน</span>'
        jj = pj.get('j', '—')
        jc = pc.get('j', '—')

        w18_req_j = rj.get('w18_req')  if rj  else None
        w18_req_c = rc_.get('w18_req') if rc_ else None
        w18_cap_j = rj.get('w18_cap')  if rj  else None
        w18_cap_c = rc_.get('w18_cap') if rc_ else None
        ratio_j   = rj.get('ratio')    if rj  else None
        ratio_c   = rc_.get('ratio')   if rc_ else None
        passed_j  = rj.get('passed')   if rj  else None
        passed_c  = rc_.get('passed')  if rc_ else None

        # ── pre-format ก่อนเข้า f-string เพื่อหลีกเลี่ยง double-brace bug ──
        w18_req_j_s = f'{w18_req_j:,.0f}' if w18_req_j is not None else '—'
        w18_req_c_s = f'{w18_req_c:,.0f}' if w18_req_c is not None else '—'
        w18_cap_j_s = f'{w18_cap_j:,.0f}' if w18_cap_j is not None else '—'
        w18_cap_c_s = f'{w18_cap_c:,.0f}' if w18_cap_c is not None else '—'
        ratio_j_s   = f'×{ratio_j:.2f}'   if ratio_j   is not None else '—'
        ratio_c_s   = f'×{ratio_c:.2f}'   if ratio_c   is not None else '—'
        kj_eff_s    = f'{kj_eff:.0f}'     if kj_eff    is not None else '—'
        kc_eff_s    = f'{kc_eff:.0f}'     if kc_eff    is not None else '—'
        kj_opt_s    = f'{kj_opt:.0f}'     if kj_opt    is not None else '—'
        kc_opt_s    = f'{kc_opt:.0f}'     if kc_opt    is not None else '—'

        kj_ok = (kj_eff >= kj_opt) if (kj_eff and kj_opt) else None
        kc_ok = (kc_eff >= kc_opt) if (kc_eff and kc_opt) else None
        dkj   = (kj_eff - kj_opt)  if (kj_eff and kj_opt) else None
        dkc   = (kc_eff - kc_opt)  if (kc_eff and kc_opt) else None

        overall_j = (passed_j and kj_ok) if (passed_j is not None and kj_ok is not None) else None
        overall_c = (passed_c and kc_ok) if (passed_c is not None and kc_ok is not None) else None

        html = f'''
        <table class="cmp-table">
          <thead>
            <tr>
              <th class="cmp-th-label">รายการ</th>
              <th class="cmp-th-j">◻ JPCP / JRCP</th>
              <th class="cmp-th-c">◻ CRCP</th>
            </tr>
          </thead>
          <tbody>

            <tr class="cmp-shdr"><td colspan="3">1 · พารามิเตอร์ออกแบบ</td></tr>
            <tr>
              <td class="cmp-td-label">f'c (cube)</td>
              <td class="cmp-td-shared" colspan="2">{fc_cube:.0f} ksc {shared}</td>
            </tr>
            <tr>
              <td class="cmp-td-label">Ec</td>
              <td class="cmp-td-shared" colspan="2">{ec_psi:,.0f} psi {shared}</td>
            </tr>
            <tr>
              <td class="cmp-td-label">Sc (ทล. lock)</td>
              <td class="cmp-td-shared" colspan="2">{SC_FIXED:.0f} psi {shared}</td>
            </tr>
            <tr>
              <td class="cmp-td-label">J — Load Transfer</td>
              <td class="cmp-td-b">{jj}</td>
              <td class="cmp-td-g">{jc}</td>
            </tr>
            <tr>
              <td class="cmp-td-label">Cd — Drainage</td>
              <td class="cmp-td-shared" colspan="2">{cd:.1f} {shared}</td>
            </tr>
            <tr>
              <td class="cmp-td-label">Pt / ΔPSI</td>
              <td class="cmp-td-shared" colspan="2">{pt:.1f} / {dpsi:.1f} {shared}</td>
            </tr>
            <tr>
              <td class="cmp-td-label">ZR / So</td>
              <td class="cmp-td-shared" colspan="2">{zr:.3f} / {so:.2f} {shared}</td>
            </tr>
            <tr>
              <td class="cmp-td-label">CBR (subgrade)</td>
              <td class="cmp-td-shared" colspan="2">{cbr:.1f} % {shared}</td>
            </tr>
            <tr>
              <td class="cmp-td-label">MR (subgrade)</td>
              <td class="cmp-td-shared" colspan="2">{MR_psi:,.0f} psi {shared}</td>
            </tr>

            <tr class="cmp-shdr"><td colspan="3">2 · ความหนาแผ่นคอนกรีต</td></tr>
            <tr>
              <td class="cmp-td-label">D แนะนำ</td>
              <td class="cmp-td-b">
                <span style="font-family:'IBM Plex Mono',monospace;font-size:20px;font-weight:700;color:#1565C0">{dj_in} in</span>
                <div style="font-size:10px;color:#78909C">{dj_cm or '—'} ซม.</div>
              </td>
              <td class="cmp-td-g">
                <span style="font-family:'IBM Plex Mono',monospace;font-size:20px;font-weight:700;color:#2E7D32">{dc_in} in</span>
                <div style="font-size:10px;color:#78909C">{dc_cm or '—'} ซม.</div>
              </td>
            </tr>
            <tr>
              <td class="cmp-td-label">W18 required (D แนะนำ)</td>
              <td class="cmp-td-b">{w18_req_j_s}</td>
              <td class="cmp-td-g">{w18_req_c_s}</td>
            </tr>
            <tr>
              <td class="cmp-td-label">W18 capacity</td>
              <td class="cmp-td-b">{w18_cap_j_s}
                {_bar(ratio_j, 'cmp-bar-b') if ratio_j else ''}
              </td>
              <td class="cmp-td-g">{w18_cap_c_s}
                {_bar(ratio_c, 'cmp-bar-g') if ratio_c else ''}
              </td>
            </tr>
            <tr>
              <td class="cmp-td-label">Ratio (cap / req)</td>
              <td class="cmp-td-b" style="font-weight:700">{ratio_j_s}</td>
              <td class="cmp-td-g" style="font-weight:700">{ratio_c_s}</td>
            </tr>

            <tr class="cmp-shdr"><td colspan="3">3 · k_opt vs k_eff</td></tr>
            <tr>
              <td class="cmp-td-label">k_eff (จาก Tab 2)</td>
              <td class="cmp-td-b">{kj_eff_s} pci</td>
              <td class="cmp-td-g">{kc_eff_s} pci</td>
            </tr>
            <tr>
              <td class="cmp-td-label">k_opt (min required)</td>
              <td class="cmp-td-b">{kj_opt_s} pci</td>
              <td class="cmp-td-g">{kc_opt_s} pci</td>
            </tr>
            <tr>
              <td class="cmp-td-label">Δk = k_eff − k_opt</td>
              <td class="cmp-td-b">{(f"{dkj:+.0f} pci ({dkj/kj_opt*100:+.1f}%)") if dkj is not None else "—"}</td>
              <td class="cmp-td-g">{(f"{dkc:+.0f} pci ({dkc/kc_opt*100:+.1f}%)") if dkc is not None else "—"}</td>
            </tr>

            <tr class="cmp-shdr"><td colspan="3">4 · ผลการตรวจสอบ</td></tr>
            <tr>
              <td class="cmp-td-label">W18 capacity ≥ W18 required</td>
              <td class="cmp-td-b">{_badge(passed_j)}</td>
              <td class="cmp-td-g">{_badge(passed_c)}</td>
            </tr>
            <tr>
              <td class="cmp-td-label">k_eff ≥ k_opt</td>
              <td class="cmp-td-b">{_badge(kj_ok)}</td>
              <td class="cmp-td-g">{_badge(kc_ok)}</td>
            </tr>
            <tr style="background:#F9FBE7">
              <td class="cmp-td-label" style="font-weight:700;color:#33691E">สรุปผล</td>
              <td style="text-align:center">{_badge(overall_j)}</td>
              <td style="text-align:center">{_badge(overall_c)}</td>
            </tr>

          </tbody>
        </table>'''

        # ── เตรียม sections สำหรับ PDF ────────────────────────
        dpsi_val = 4.5 - pt
        pdf_sections = [
            ('1 · พารามิเตอร์ออกแบบ', [
                {'label': "f'c (cube)",    'val_j': f'{fc_cube:.0f} ksc',         'shared': True},
                {'label': 'Ec',            'val_j': f'{ec_psi:,.0f} psi',          'shared': True},
                {'label': 'Sc (ทล. lock)', 'val_j': f'{SC_FIXED:.0f} psi',         'shared': True},
                {'label': 'J',             'val_j': str(jj), 'val_c': str(jc)},
                {'label': 'Cd',            'val_j': f'{cd:.1f}',                   'shared': True},
                {'label': 'Pt / DPSI',     'val_j': f'{pt:.1f} / {dpsi_val:.1f}', 'shared': True},
                {'label': 'ZR / So',       'val_j': f'{zr:.3f} / {so:.2f}',       'shared': True},
                # ── เพิ่ม CBR และ MR ──────────────────────────
                {'label': 'CBR',           'val_j': f'{cbr:.1f} %',               'shared': True},
                {'label': 'MR (subgrade)', 'val_j': f'{MR_psi:,.0f} psi',         'shared': True},
            ]),
            ('2 · ความหนาแผ่นคอนกรีต', [
                {'label': 'D แนะนำ',
                 'val_j': f'{dj_in} in ({dj_cm} cm)' if dj_cm else '-',
                 'val_c': f'{dc_in} in ({dc_cm} cm)' if dc_cm else '-', 'bold': True},
                {'label': 'W18 required (D แนะนำ)',
                 'val_j': f'{w18_req_j:,.0f}' if w18_req_j else '-',
                 'val_c': f'{w18_req_c:,.0f}' if w18_req_c else '-'},
                {'label': 'W18 capacity',
                 'val_j': f'{w18_cap_j:,.0f}' if w18_cap_j else '-',
                 'val_c': f'{w18_cap_c:,.0f}' if w18_cap_c else '-'},
                {'label': 'Ratio (cap/req)',
                 'val_j': f'x{ratio_j:.2f}' if ratio_j else '-',
                 'val_c': f'x{ratio_c:.2f}' if ratio_c else '-', 'bold': True},
            ]),
            ('3 · k_opt vs k_eff', [
                {'label': 'k_eff (Tab 2)',
                 'val_j': f'{kj_eff:.0f} pci' if kj_eff else '-',
                 'val_c': f'{kc_eff:.0f} pci' if kc_eff else '-'},
                {'label': 'k_opt (min required)',
                 'val_j': f'{kj_opt:.0f} pci' if kj_opt else '-',
                 'val_c': f'{kc_opt:.0f} pci' if kc_opt else '-'},
                {'label': 'Dk = k_eff - k_opt',
                 'val_j': f'{dkj:+.0f} pci ({dkj/kj_opt*100:+.1f}%)' if dkj else '-',
                 'val_c': f'{dkc:+.0f} pci ({dkc/kc_opt*100:+.1f}%)' if dkc else '-'},
            ]),
            ('4 · ผลการตรวจสอบ', [
                {'label': 'W18 cap >= W18 req',
                 'val_j': 'ผ่าน' if passed_j else 'ไม่ผ่าน',
                 'val_c': 'ผ่าน' if passed_c else 'ไม่ผ่าน', 'bold': True},
                {'label': 'k_eff >= k_opt',
                 'val_j': 'ผ่าน' if kj_ok else 'ไม่ผ่าน',
                 'val_c': 'ผ่าน' if kc_ok else 'ไม่ผ่าน', 'bold': True},
                {'label': 'สรุปผล',
                 'val_j': 'ผ่าน' if overall_j else 'ไม่ผ่าน',
                 'val_c': 'ผ่าน' if overall_c else 'ไม่ผ่าน',
                 'bold': True, 'shade': True},
            ]),
        ]
        st.session_state['_pdf_sections'] = pdf_sections
        st.session_state['_pdf_layers_j'] = layers_j
        st.session_state['_pdf_layers_c'] = layers_c
        st.session_state['_pdf_dj_cm']    = dj_cm
        st.session_state['_pdf_dc_cm']    = dc_cm

        # render ผ่าน components.html
        n_rows   = 9 + 4 + 3 + 3   # section 1 เพิ่ม 2 rows (CBR + MR)
        height   = n_rows * 32 + 140
        full_html = f'''<!DOCTYPE html><html><head>
        <meta charset="utf-8">
        <link href="https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;600&family=Sarabun:wght@300;400;600&display=swap" rel="stylesheet">
        {TABLE_CSS}
        <style>body{{margin:0;padding:4px 0;font-family:"Sarabun",sans-serif;}}</style>
        </head><body>{html}</body></html>'''
        components.html(full_html, height=height, scrolling=False)

    # ════════════════════════════════════════════════════════
    # Card 2 — Layer Structure
    # ════════════════════════════════════════════════════════
    if layers_j or layers_c:
        with st.container(border=True):
            st.markdown('<div class="rp-card-title">🏗️ โครงสร้างชั้นทาง — ความหนาแต่ละชั้น (ซม.)</div>',
                        unsafe_allow_html=True)

            names_j   = [l['name'] for l in layers_j]
            names_c   = [l['name'] for l in layers_c]
            all_names = list(dict.fromkeys(names_j + names_c))

            def _get_thick(layers, name):
                for l in layers:
                    if l['name'] == name:
                        return l['thickness_cm']
                return None

            rows_html = ''
            rows_html += f'''
            <tr style="background:#EEF2F7">
              <td class="cmp-ltd-no" style="font-weight:700">0</td>
              <td class="cmp-ltd-mat" style="font-weight:700;color:#1565C0">แผ่นคอนกรีต (D)</td>
              <td class="cmp-ltd-j" style="font-size:14px">{dj_cm or "—"}</td>
              <td class="cmp-ltd-c" style="font-size:14px">{dc_cm or "—"}</td>
            </tr>'''

            tot_j = dj_cm or 0
            tot_c = dc_cm or 0
            for i, name in enumerate(all_names, start=1):
                tj = _get_thick(layers_j, name)
                tc = _get_thick(layers_c, name)
                if tj: tot_j += tj
                if tc: tot_c += tc
                td_j = f'<td class="cmp-ltd-j">{tj}</td>' if tj else '<td class="cmp-ltd-na">—</td>'
                td_c = f'<td class="cmp-ltd-c">{tc}</td>' if tc else '<td class="cmp-ltd-na">—</td>'
                rows_html += f'''
                <tr>
                  <td class="cmp-ltd-no">{i}</td>
                  <td class="cmp-ltd-mat">{name}</td>
                  {td_j}{td_c}
                </tr>'''

            rows_html += f'''
            <tr class="cmp-total">
              <td class="cmp-ltd-no"></td>
              <td class="cmp-ltd-mat">รวมทั้งหมด (รวมคอนกรีต)</td>
              <td class="cmp-ltd-j" style="font-size:13px">{tot_j} ซม.</td>
              <td class="cmp-ltd-c" style="font-size:13px">{tot_c} ซม.</td>
            </tr>'''

            n_layer_rows = len(all_names) + 2
            layer_height = n_layer_rows * 30 + 60
            layer_full = f'''<!DOCTYPE html><html><head>
            <style>body{{margin:0;padding:0;font-family:"Sarabun",sans-serif;}}</style>
            </head><body>
            {TABLE_CSS}
            <table class="cmp-layer-table">
              <thead><tr>
                <th class="cmp-lth-no">#</th>
                <th class="cmp-lth-mat">วัสดุ</th>
                <th class="cmp-lth-j">◻ JPCP (ซม.)</th>
                <th class="cmp-lth-c">◻ CRCP (ซม.)</th>
              </tr></thead>
              <tbody>{rows_html}</tbody>
            </table></body></html>'''
            components.html(layer_full, height=layer_height, scrolling=False)

    # ── PDF Export button ─────────────────────────────────────
    st.markdown('<div style="height:8px"></div>', unsafe_allow_html=True)
    with st.container(border=True):
        st.markdown('<div class="rp-card-title">📄 Export PDF Summary</div>',
                    unsafe_allow_html=True)
        proj_name = st.session_state.get('project_name', '') or '(ไม่ระบุชื่อโครงการ)'
        date_str  = datetime.now().strftime('%d/%m/%Y %H:%M')
        secs = st.session_state.get('_pdf_sections')
        lj   = st.session_state.get('_pdf_layers_j', [])
        lc   = st.session_state.get('_pdf_layers_c', [])
        dj   = st.session_state.get('_pdf_dj_cm')
        dc   = st.session_state.get('_pdf_dc_cm')
        if secs:
            try:
                pdf_buf = _create_pdf_summary(proj_name, date_str, secs, lj, lc, dj, dc)
                if pdf_buf:
                    proj_slug = proj_name.replace(' ', '_')[:20] \
                        if proj_name != '(ไม่ระบุชื่อโครงการ)' else 'NoName'
                    fname = f'Summary_{proj_slug}_{datetime.now().strftime("%Y%m%d_%H%M")}.pdf'
                    st.download_button(
                        '📥 Download PDF Summary', pdf_buf, fname,
                        'application/pdf', key='dl_pdf_summary',
                        use_container_width=True)
                else:
                    st.error('❌ ไม่พบ fpdf2 — กรุณาเพิ่ม fpdf2 ใน requirements.txt')
            except Exception as e:
                st.error(f'❌ สร้าง PDF ไม่สำเร็จ: {e}')
        else:
            st.markdown(
                '<div class="rp-status-info">ℹ️ คำนวณ Design ทั้ง JPCP และ CRCP ก่อนเพื่อ export PDF</div>',
                unsafe_allow_html=True)
