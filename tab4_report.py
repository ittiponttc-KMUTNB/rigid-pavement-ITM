"""
tab4_report.py — Tab 4: Report
Rigid Pavement Design V7
สร้าง Word report ฉบับสมบูรณ์ — port จาก V6
"""
import streamlit as st
from io import BytesIO
from datetime import datetime
import matplotlib.pyplot as plt

from engine import (
    plot_structure, fig_to_bytes, get_zr,
    LAYER_NAMES_EN, mr_from_cbr,
)

# ============================================================
# Default texts
# ============================================================
DEFAULT_INTRO = (
    'การออกแบบความหนาแผ่นคอนกรีตตามแนวทางของ AASHTO 1993 จำเป็นต้องอาศัยสมเหตุสมผลที่'
    'พัฒนามาจากผลการทดสอบ AASHO Road Test ซึ่งสะท้อนพฤติกรรมการรับน้ำหนักและการเสื่อมสภาพของแผ่น'
    'คอนกรีตภายใต้สภาพการใช้งานจริง สมการดังกล่าวรวมปัจจัยสำคัญหลายด้าน ทั้งด้านปริมาณจราจร '
    'ความน่าเชื่อถือของการออกแบบ คุณสมบัติวัสดุ และสภาพชั้นรองรับ เพื่อให้สามารถประเมินความหนา'
    'ที่เหมาะสมสำหรับรองรับปริมาณจราจรตลอดอายุโครงการได้อย่างแม่นยำ '
    'สมการหลักที่ใช้ในการออกแบบความหนาถนนคอนกรีตตาม AASHTO 1993 มีดังนี้'
)
DEFAULT_PAVEMENT_DESC = (
    'โดยมาตรฐานการออกแบบตามวิธี AASHTO 1993 ได้แบ่งโครงสร้างทางคอนกรีตออกเป็นหลายรูปแบบตาม'
    'ลักษณะการควบคุมความแตกร้าวและการถ่ายแรงระหว่างแผ่นคอนกรีต แต่ละประเภทมีแนวคิดการออกแบบ'
    'และยุทธวิธีดำเนินโครงการก่อสร้างที่แตกต่างกัน โครงสร้างทางคอนกรีต 3 ประเภทหลักสำหรับการคำนวณ'
    ' ได้แก่ Jointed Plain Concrete Pavement (JPCP), Jointed Reinforced Concrete Pavement (JRCP) '
    'และ Continuously Reinforced Concrete Pavement (CRCP)'
)
DEFAULT_SUMMARY = (
    'จากการคำนวณตามวิธีของ AASHTO 1993 ผิวทางคอนกรีต (Concrete Pavement) สามารถสรุปรูปแบบของ'
    'โครงสร้างชั้นทางที่ออกแบบได้ดังแสดงในตารางและรูปด้านล่าง'
)


# ============================================================
# Word Report Engine (port จาก V6)
# ============================================================
def _get_font(): return 'TH SarabunPSK'

def _fmt_name(name): return LAYER_NAMES_EN.get(name, name)

def _build_report(
    # หัวข้อ/เลขรูป
    sec_prefix, fig_prefix, fig_start,
    intro_text, pavement_desc, summary_text,
    # โครงการ
    proj_name, calc_date,
    # JPCP
    inc_jpcp, jpcp_layers, jpcp_params, jpcp_rows,
    jpcp_fig33, jpcp_fig34,
    # CRCP
    inc_crcp, crcp_layers, crcp_params, crcp_rows,
    crcp_fig33, crcp_fig34,
    # options
    inc_summary,
):
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return None

    TH  = _get_font()
    EQ  = 'Times New Roman'
    TS  = Pt(15)
    EQS = Pt(11)
    HBG = 'BDD7EE'
    SBG = 'FFF2CC'
    PBG = 'CCFFCC'
    FBG = 'FFCCCC'
    SEL = 'FFFFAA'

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = TH; style.font.size = TS
    sec = doc.sections[0]
    sec.page_width=Cm(21); sec.page_height=Cm(29.7)
    sec.left_margin=sec.right_margin=Cm(2.5)
    sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.0)

    fig_counter = [fig_start]
    def nfig():
        n = fig_counter[0]; fig_counter[0] += 1; return n

    # ── helpers ──────────────────────────────────────────────
    def _cell_fmt(cell, bg=None):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        tcMar=OxmlElement('w:tcMar')
        for side in ['top','bottom','left','right']:
            m=OxmlElement(f'w:{side}')
            m.set(qn('w:w'),'80'); m.set(qn('w:type'),'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)
        if bg:
            shd=OxmlElement('w:shd')
            shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
            shd.set(qn('w:fill'),bg); tcPr.append(shd)

    def _set_col_w(row, widths):
        for i,cell in enumerate(row.cells):
            tc=cell._tc; tcPr=tc.get_or_add_tcPr()
            tcW=OxmlElement('w:tcW')
            tcW.set(qn('w:w'),str(widths[i])); tcW.set(qn('w:type'),'dxa')
            tcPr.append(tcW)

    def _sc(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, bg=None, size=None):
        cell.text=''
        p=cell.paragraphs[0]; p.alignment=align
        run=p.add_run(text)
        run.font.name=TH; run.font.size=size or TS; run.bold=bold
        _cell_fmt(cell, bg)

    def _set_valign(cell, val='center'):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        vA=OxmlElement('w:vAlign'); vA.set(qn('w:val'),val); tcPr.append(vA)

    def _set_vmerge(cell, restart=False):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        vM=OxmlElement('w:vMerge')
        if restart: vM.set(qn('w:val'),'restart')
        tcPr.append(vM)

    def _set_cw(cell, w):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        tcW=OxmlElement('w:tcW')
        tcW.set(qn('w:w'),str(w)); tcW.set(qn('w:type'),'dxa')
        tcPr.append(tcW)

    def _cell_margin(cell, mar=80):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        tcMar=OxmlElement('w:tcMar')
        for side in ['top','bottom','left','right']:
            m=OxmlElement(f'w:{side}')
            m.set(qn('w:w'),str(mar)); m.set(qn('w:type'),'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)

    def _add_heading(text, level=1):
        p=doc.add_paragraph()
        run=p.add_run(text)
        run.font.name=TH; run.font.size=TS
        run.bold=True; run.underline=(level<=2)
        return p

    def _add_para(text, bold=False, indent_cm=0):
        p=doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        run=p.add_run(text)
        run.font.name=TH; run.font.size=TS; run.bold=bold
        if indent_cm>0: p.paragraph_format.left_indent=Cm(indent_cm)
        return p

    def _add_fig_caption(text):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(text)
        run.font.name=TH; run.font.size=TS; run.bold=True; run.underline=True

    def _eq_run(p, text, sub=False, sup=False, bold=False):
        run=p.add_run(text)
        run.font.name=EQ; run.font.size=EQS; run.bold=bold
        if sub or sup:
            rPr=run._r.get_or_add_rPr()
            va=OxmlElement('w:vertAlign')
            va.set(qn('w:val'),'subscript' if sub else 'superscript')
            rPr.append(va)
        return run

    def _th_run(p, text, bold=False):
        run=p.add_run(text)
        run.font.name=TH; run.font.size=TS; run.bold=bold
        return run

    def _eq_line(indent_cm=1.5):
        p=doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent=Cm(indent_cm)
        p.paragraph_format.space_after=Pt(2)
        return p

    def _sec_num(base, sub=None):
        if sub is None: return base
        return f'{base}.{sub}'

    # ── หน้าปก ───────────────────────────────────────────────
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('รายการคำนวณออกแบบ\nผิวทางคอนกรีต')
    r.font.name=TH; r.font.size=Pt(20); r.bold=True

    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('ตามวิธี AASHTO 1993')
    r.font.name=TH; r.font.size=Pt(16)

    if proj_name:
        doc.add_paragraph()
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(f'โครงการ: {proj_name}')
        r.font.name=TH; r.font.size=TS

    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f'วันที่: {calc_date}')
    r.font.name=TH; r.font.size=TS

    doc.add_page_break()

    # ── หัวข้อหลัก + บทเกริ่นนำ + สมการ ─────────────────────
    _add_heading(f'{sec_prefix}  การออกแบบผิวทางคอนกรีต', level=1)
    _add_para(intro_text)
    doc.add_paragraph()

    # สมการ AASHTO 1993
    p1=_eq_line()
    _eq_run(p1,'log'); _eq_run(p1,'10',sub=True)
    _eq_run(p1,'(W');  _eq_run(p1,'18',sub=True)
    _eq_run(p1,') = Z'); _eq_run(p1,'R',sub=True)
    _eq_run(p1,' \u00d7 S'); _eq_run(p1,'o',sub=True)
    _eq_run(p1,' + 7.35 \u00d7 log'); _eq_run(p1,'10',sub=True)
    _eq_run(p1,'(D+1) \u2212 0.06')

    p2=_eq_line()
    _eq_run(p2,'        + log'); _eq_run(p2,'10',sub=True)
    _eq_run(p2,'(\u0394PSI / (4.5 \u2212 1.5)) / (1 + 1.624\u00d710')
    _eq_run(p2,'7',sup=True)
    _eq_run(p2,' / (D+1)'); _eq_run(p2,'8.46',sup=True); _eq_run(p2,')')

    p3=_eq_line()
    _eq_run(p3,'        + (4.22 \u2212 0.32\u00d7P'); _eq_run(p3,'t',sub=True)
    _eq_run(p3,') \u00d7 log'); _eq_run(p3,'10',sub=True)
    _eq_run(p3,' [(S'); _eq_run(p3,'c',sub=True)
    _eq_run(p3,'\u00d7C'); _eq_run(p3,'d',sub=True)
    _eq_run(p3,'\u00d7(D'); _eq_run(p3,'0.75',sup=True)
    _eq_run(p3,'\u22121.132)) / (215.63\u00d7J\u00d7(D'); _eq_run(p3,'0.75',sup=True)
    _eq_run(p3,' \u2212 18.42 / (E'); _eq_run(p3,'c',sub=True)
    _eq_run(p3,'/k)'); _eq_run(p3,'0.25',sup=True); _eq_run(p3,')]')

    doc.add_paragraph()
    p_by=doc.add_paragraph(); _th_run(p_by,'โดยที่:')

    cw_sym=[1396,6281,1395]
    tsym=doc.add_table(rows=1,cols=3)
    tsym.style='Table Grid'; tsym.alignment=WD_TABLE_ALIGNMENT.LEFT
    hdr=tsym.rows[0]; _set_col_w(hdr,cw_sym)
    for i,h in enumerate(['สัญลักษณ์','ความหมาย','หน่วย']):
        _sc(hdr.cells[i],h,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,bg=HBG)
    syms=[
        ('W\u2081\u2088','จำนวนแกนเดี่ยว 18 kip ที่รองรับได้','ESALs'),
        ('ZR','Standard Normal Deviate ที่ความเชื่อมั่น R','-'),
        ('So','Overall Standard Deviation','-'),
        ('D','ความหนาแผ่นคอนกรีต','นิ้ว'),
        ('ΔPSI','การสูญเสีย Serviceability (4.5 − Pt)','-'),
        ('Pt','Terminal Serviceability Index','-'),
        ('Sc','Modulus of Rupture ของคอนกรีต','psi'),
        ('Cd','Drainage Coefficient','-'),
        ('J','Load Transfer Coefficient','-'),
        ('Ec','Modulus of Elasticity ของคอนกรีต','psi'),
        ('k','Modulus of Subgrade Reaction','pci'),
    ]
    for sym,meaning,unit in syms:
        row=tsym.add_row(); _set_col_w(row,cw_sym)
        _sc(row.cells[0],sym,align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(row.cells[1],meaning)
        _sc(row.cells[2],unit,align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    _add_para(pavement_desc)
    doc.add_paragraph()

    # ── helper: layer table ───────────────────────────────────
    def _layer_table(layers, d_cm, ptype, cbr):
        cw=[756,4536,1728,2052]
        tbl=doc.add_table(rows=1,cols=4)
        tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.LEFT
        hdr=tbl.rows[0]; _set_col_w(hdr,cw)
        for i,h in enumerate(['ลำดับ','ชนิดวัสดุ','ความหนา (ซม.)','Modulus E (MPa)']):
            _sc(hdr.cells[i],h,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,bg=HBG)
        row=tbl.add_row(); _set_col_w(row,cw)
        _sc(row.cells[0],'1',align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(row.cells[1],f'ผิวทางคอนกรีต {ptype}')
        _sc(row.cells[2],str(d_cm),align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(row.cells[3],'-',align=WD_ALIGN_PARAGRAPH.CENTER)
        rn=2
        for layer in layers:
            t=layer.get('thickness_cm',0)
            if t<=0: continue
            e=layer.get('E_MPa',0)
            row=tbl.add_row(); _set_col_w(row,cw)
            _sc(row.cells[0],str(rn),align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(row.cells[1],_fmt_name(layer.get('name','')))
            _sc(row.cells[2],str(t),align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(row.cells[3],f'{e:,}' if e>0 else '-',align=WD_ALIGN_PARAGRAPH.CENTER)
            rn+=1
        mr_psi=int(1500*cbr if cbr<10 else 1000+555*cbr)
        mr_mpa=round(mr_psi/145.038)
        row=tbl.add_row(); _set_col_w(row,cw)
        _sc(row.cells[0],str(rn),align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(row.cells[1],'ดินคันทาง')
        _sc(row.cells[2],f'CBR \u2265 {cbr:.1f} %',align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(row.cells[3],f'{mr_mpa:,} ({mr_psi:,} psi)',align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

    # ── helper: ESB calculation ───────────────────────────────
    def _esb_section(layers):
        valid=[l for l in layers if l.get('thickness_cm',0)>0 and l.get('E_MPa',0)>0]
        if not valid: return
        _add_heading('การคำนวณ Subbase Elastic Modulus (ESB)',level=2)
        p_d=doc.add_paragraph(); p_d.alignment=WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        _th_run(p_d,'ค่า Subbase Elastic Modulus (E')
        r=p_d.add_run('SB'); r.font.name=EQ; r.font.size=EQS
        rPr=r._r.get_or_add_rPr(); va=OxmlElement('w:vertAlign')
        va.set(qn('w:val'),'subscript'); rPr.append(va)
        _th_run(p_d,') คำนวณจากโมดูลัสเทียบเท่าของชั้นวัสดุรองพื้นทาง โดยใช้สมการดังนี้')
        peq=_eq_line(2.0)
        _eq_run(peq,'E'); _eq_run(peq,'SB',sub=True)
        _eq_run(peq,'  =  [  \u03a3 ( h')
        _eq_run(peq,'i',sub=True); _eq_run(peq,'  \u00d7  E')
        _eq_run(peq,'i',sub=True); _eq_run(peq,'1/3',sup=True)
        _eq_run(peq,' )  /  \u03a3 h'); _eq_run(peq,'i',sub=True)
        _eq_run(peq,'  ]'); _eq_run(peq,'3',sup=True)
        doc.add_paragraph()
        _th_run(doc.add_paragraph(),'การคำนวณแสดงในตารางดังนี้')
        cw2=[570,2900,1400,1400,1400,1400]
        tbl2=doc.add_table(rows=1,cols=6)
        tbl2.style='Table Grid'; tbl2.alignment=WD_TABLE_ALIGNMENT.LEFT
        hdr2=tbl2.rows[0]; _set_col_w(hdr2,cw2)
        def _hc(cell, parts, bg=HBG):
            cell.text=''; p=cell.paragraphs[0]
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER; _cell_fmt(cell,bg)
            for text,font,fsize,sup,sub in parts:
                r=p.add_run(text); r.font.name=font; r.font.size=fsize; r.bold=True
                if sup:
                    rP=r._r.get_or_add_rPr(); v=OxmlElement('w:vertAlign')
                    v.set(qn('w:val'),'superscript'); rP.append(v)
                if sub:
                    rP=r._r.get_or_add_rPr(); v=OxmlElement('w:vertAlign')
                    v.set(qn('w:val'),'subscript'); rP.append(v)
        _hc(hdr2.cells[0],[('ลำดับ',TH,TS,False,False)])
        _hc(hdr2.cells[1],[('ชั้นวัสดุ',TH,TS,False,False)])
        _hc(hdr2.cells[2],[('h',EQ,EQS,False,False),('i',EQ,EQS,False,True),(' (ซม.)',TH,TS,False,False)])
        _hc(hdr2.cells[3],[('E',EQ,EQS,False,False),('i',EQ,EQS,False,True),(' (MPa)',TH,TS,False,False)])
        _hc(hdr2.cells[4],[('E',EQ,EQS,False,False),('i',EQ,EQS,False,True),('1/3',EQ,EQS,True,False)])
        _hc(hdr2.cells[5],[('h',EQ,EQS,False,False),('i',EQ,EQS,False,True),
                            (' \u00d7 E',EQ,EQS,False,False),('i',EQ,EQS,False,True),('1/3',EQ,EQS,True,False)])
        def _td(cell,text,bold=False,align=WD_ALIGN_PARAGRAPH.CENTER,bg=None):
            cell.text=''; p=cell.paragraphs[0]; p.alignment=align
            r=p.add_run(text); r.font.name=TH; r.font.size=TS; r.bold=bold
            _cell_fmt(cell,bg)
        sh=0.0; shE=0.0
        for idx,layer in enumerate(valid,1):
            h=layer['thickness_cm']; E=layer['E_MPa']
            E13=E**(1/3); hE=h*E13; sh+=h; shE+=hE
            r2=tbl2.add_row(); _set_col_w(r2,cw2)
            _td(r2.cells[0],str(idx))
            _td(r2.cells[1],_fmt_name(layer.get('name','')),align=WD_ALIGN_PARAGRAPH.LEFT)
            _td(r2.cells[2],f'{h:,}')
            _td(r2.cells[3],f'{E:,}')
            _td(r2.cells[4],f'{E13:.4f}')
            _td(r2.cells[5],f'{hE:,.2f}')
        rs=tbl2.add_row(); _set_col_w(rs,cw2)
        _td(rs.cells[0],'',bg=SBG)
        _td(rs.cells[1],'รวม (\u03a3)',bold=True,align=WD_ALIGN_PARAGRAPH.RIGHT,bg=SBG)
        _td(rs.cells[2],f'{sh:.0f}',bold=True,bg=SBG)
        _td(rs.cells[3],'',bg=SBG); _td(rs.cells[4],'',bg=SBG)
        _td(rs.cells[5],f'{shE:,.2f}',bold=True,bg=SBG)
        if sh>0:
            esb_mpa=(shE/sh)**3; esb_psi=esb_mpa*145.038
            doc.add_paragraph()
            pr1=_eq_line(1.5); _th_run(pr1,'แทนค่า  ')
            _eq_run(pr1,'E'); _eq_run(pr1,'SB',sub=True)
            _eq_run(pr1,f'  =  [ {shE:,.2f} / {sh:.0f} ]'); _eq_run(pr1,'3',sup=True)
            pr2=_eq_line(1.5); _th_run(pr2,'ดังนั้น  ')
            _eq_run(pr2,'E'); _eq_run(pr2,'SB',sub=True)
            _eq_run(pr2,f'  =  {esb_mpa:,.2f}'); _th_run(pr2,'  MPa')
            _eq_run(pr2,f'  =  {esb_psi:,.0f}'); _th_run(pr2,'  psi',bold=True)
        doc.add_paragraph()

    # ── helper: k-value section ───────────────────────────────
    def _kvalue_section(params, fig33, fig34, fig_n_k):
        cw_k=[5772,1924,1376]
        _add_para('ขั้นตอนที่ 1: หาค่า Composite Modulus of Subgrade Reaction (k\u221e)',bold=True)
        tk1=doc.add_table(rows=1,cols=3)
        tk1.style='Table Grid'; tk1.alignment=WD_TABLE_ALIGNMENT.LEFT
        hdr=tk1.rows[0]; _set_col_w(hdr,cw_k)
        for i,h in enumerate(['พารามิเตอร์','ค่า','หน่วย']):
            _sc(hdr.cells[i],h,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,bg=HBG)
        for p_n,val,unit in [
            ('Roadbed Soil Resilient Modulus (MR)',f"{params.get('MR_psi',0):,.0f}",'psi'),
            ('Subbase Elastic Modulus (ESB)',      f"{params.get('ESB_psi',0):,.0f}",'psi'),
            ('Subbase Thickness (DSB)',             f"{params.get('DSB_in',0):.1f}",'inches'),
            ('Composite Modulus k\u221e',           f"{params.get('k_inf',0):,.0f}",'pci'),
        ]:
            row=tk1.add_row(); _set_col_w(row,cw_k)
            _sc(row.cells[0],p_n)
            _sc(row.cells[1],val,align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(row.cells[2],unit,align=WD_ALIGN_PARAGRAPH.CENTER)
        if fig33:
            doc.add_paragraph()
            p_img=doc.add_paragraph(); p_img.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(BytesIO(fig33),width=Inches(5.0))
            _add_fig_caption(f'รูปที่ {fig_prefix}{fig_n_k}  ค่า Composite Modulus of Subgrade Reaction, k\u221e (pci)')
        doc.add_paragraph()
        ls=params.get('ls',0.0)
        _add_para('ขั้นตอนที่ 2: ปรับแก้ค่า Loss of Support (LS)',bold=True)
        tk2=doc.add_table(rows=1,cols=3)
        tk2.style='Table Grid'; tk2.alignment=WD_TABLE_ALIGNMENT.LEFT
        hdr=tk2.rows[0]; _set_col_w(hdr,cw_k)
        for i,h in enumerate(['พารามิเตอร์','ค่า','หน่วย']):
            _sc(hdr.cells[i],h,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,bg=HBG)
        for p_n,val,unit in [
            ('Effective Modulus k\u221e (จาก Step 1)',f"{params.get('k_inf',0):,.0f}",'pci'),
            ('Loss of Support Factor (LS)',            f"{ls:.1f}",'-'),
            ('Corrected Modulus k (ที่ใช้ออกแบบ)',   f"{params.get('k_eff',0):,.0f}",'pci'),
        ]:
            row=tk2.add_row(); _set_col_w(row,cw_k)
            _sc(row.cells[0],p_n)
            _sc(row.cells[1],val,align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(row.cells[2],unit,align=WD_ALIGN_PARAGRAPH.CENTER)
        if fig34 and ls>0:
            doc.add_paragraph()
            p_img=doc.add_paragraph(); p_img.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(BytesIO(fig34),width=Inches(5.0))
            _add_fig_caption(f'รูปที่ {fig_prefix}{fig_n_k+1}  การปรับแก้ค่า Modulus of Subgrade Reaction ประสิทธิผล เนื่องจากการสูญเสียฐานรองรับ')
        doc.add_paragraph()

    # ── helper: design result ─────────────────────────────────
    def _design_result(params, rows, sel_d_cm):
        # ข้อมูลนำเข้า
        p_l=doc.add_paragraph()
        r=p_l.add_run('ข้อมูลนำเข้าการออกแบบ:')
        r.font.name=TH; r.font.size=TS; r.bold=True; r.underline=True
        cw_in=[3923,1471,2207,1471]
        tin=doc.add_table(rows=1,cols=4)
        tin.style='Table Grid'; tin.alignment=WD_TABLE_ALIGNMENT.LEFT
        hdr=tin.rows[0]; _set_col_w(hdr,cw_in)
        for i,h in enumerate(['พารามิเตอร์','สัญลักษณ์','ค่า','หน่วย']):
            _sc(hdr.cells[i],h,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,bg=HBG)
        dpsi=params.get('dpsi',4.5-params.get('pt',2.0))
        zr=params.get('ZR',get_zr(params.get('R',90)))
        ls=params.get('ls',0.0)
        in_rows=[
            ('ESAL ออกแบบ','W\u2081\u2088',f"{params['w18']:,.0f}",'ESALs'),
            ('Terminal Serviceability','Pt',f"{params['pt']:.1f}",'-'),
            ('การสูญเสีย Serviceability','\u0394PSI',f"{dpsi:.1f}",'-'),
            ('Reliability','R',f"{params['R']:.0f}",'%'),
            ('Standard Normal Deviate','ZR',f"{zr:.3f}",'-'),
            ('Standard Deviation','So',f"{params['so']:.2f}",'-'),
            ('Modulus of Subgrade Reaction','k_eff',f"{params['k_eff']:,.0f}",'pci'),
            ('Loss of Support','LS',f"{ls:.1f}",'-'),
            ('กำลังคอนกรีต',"f'c",f"{params['fc_cube']:.0f} Cube",'ksc'),
            ('Modulus of Elasticity','Ec',f"{params['ec']:,.0f}",'psi'),
            ('Modulus of Rupture','Sc',f"{params['sc']:.0f}",'psi'),
            ('Load Transfer Coefficient','J',f"{params['j']:.1f}",'-'),
            ('Drainage Coefficient','Cd',f"{params['cd']:.2f}",'-'),
        ]
        for rd in in_rows:
            row=tin.add_row(); _set_col_w(row,cw_in)
            _sc(row.cells[0],rd[0])
            _sc(row.cells[1],rd[1],align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(row.cells[2],rd[2],align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(row.cells[3],rd[3],align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()
        # ตารางผล
        p_l2=doc.add_paragraph()
        r=p_l2.add_run('ผลการตรวจสอบความหนาแผ่นคอนกรีต:')
        r.font.name=TH; r.font.size=TS; r.bold=True; r.underline=True
        cw_res=[1188,1188,1620,2052,1512,1512]
        tres=doc.add_table(rows=1,cols=6)
        tres.style='Table Grid'; tres.alignment=WD_TABLE_ALIGNMENT.LEFT
        hdr2=tres.rows[0]; _set_col_w(hdr2,cw_res)
        for i,h in enumerate(['D (ซม.)','D (นิ้ว)','log\u2081\u2080(W\u2081\u2088)',
                               'W\u2081\u2088 รองรับได้','อัตราส่วน','ผล']):
            _sc(hdr2.cells[i],h,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,bg=HBG)
        for rv in rows:
            is_sel=(rv['d_cm']==sel_d_cm)
            bg_row=SEL if is_sel else None
            bg_res=PBG if rv['passed'] else FBG
            row2=tres.add_row(); _set_col_w(row2,cw_res)
            _sc(row2.cells[0],f"{rv['d_cm']:.0f}",bold=is_sel,align=WD_ALIGN_PARAGRAPH.CENTER,bg=bg_row)
            _sc(row2.cells[1],f"{rv['d_inch']:.0f}",align=WD_ALIGN_PARAGRAPH.CENTER,bg=bg_row)
            _sc(row2.cells[2],f"{rv['log_w18']:.4f}",align=WD_ALIGN_PARAGRAPH.CENTER,bg=bg_row)
            _sc(row2.cells[3],f"{rv['w18_cap']:,.0f}",align=WD_ALIGN_PARAGRAPH.CENTER,bg=bg_row)
            _sc(row2.cells[4],f"{rv['ratio']:.2f}",align=WD_ALIGN_PARAGRAPH.CENTER,bg=bg_row)
            _sc(row2.cells[5],'ผ่าน \u2713' if rv['passed'] else 'ไม่ผ่าน \u2717',
                align=WD_ALIGN_PARAGRAPH.CENTER,bg=bg_res)
        doc.add_paragraph()
        # สรุป
        p_l3=doc.add_paragraph()
        r=p_l3.add_run('สรุปผลการออกแบบ:')
        r.font.name=TH; r.font.size=TS; r.bold=True; r.underline=True
        sel_row=next((rv for rv in rows if rv['d_cm']==sel_d_cm),None)
        w18_cap=sel_row['w18_cap'] if sel_row else 0
        w18_req_s=sel_row.get('w18_req',params['w18']) if sel_row else params['w18']
        passed=sel_row['passed'] if sel_row else False
        ratio=sel_row['ratio']   if sel_row else 0
        for item in [
            f"ความหนาที่เลือก : {sel_d_cm:.0f} ซม. ({round(sel_d_cm/2.54):.0f} นิ้ว)",
            f"ESAL ที่ต้องการ  : {w18_req_s:,.0f} ESALs",
            f"ESAL ที่รองรับได้ : {w18_cap:,.0f} ESALs",
            f"อัตราส่วน        : {ratio:.2f}",
            f"ผลการตรวจสอบ  : {'✅ ผ่านเกณฑ์' if passed else '❌ ไม่ผ่านเกณฑ์'}",
        ]:
            p=doc.add_paragraph()
            p.paragraph_format.left_indent=Pt(36)
            run=p.add_run(item); run.font.name=TH; run.font.size=TS
        doc.add_paragraph()

    # ── helper: summary table (รูปตัดขวาง + ตาราง 3 คอลัมน์) ─
    def _summary_table(layers, d_cm, ptype, cbr, fig_caption_text):
        valid=[l for l in layers if l.get('thickness_cm',0)>0]
        data_rows=[{'thick':str(d_cm),'material':f'ผิวทางคอนกรีต\n{ptype}'}]
        for layer in valid:
            data_rows.append({'thick':str(layer.get('thickness_cm',0)),
                              'material':_fmt_name(layer.get('name',''))})
        mr_psi=int(1500*cbr if cbr<10 else 1000+555*cbr)
        data_rows.append({'thick':'Existing',
                          'material':f'Earth Embankment\nor Subgrade, CBR\u2265\n{cbr:.0f} %'})

        # สร้างรูปตัดขวาง
        fig=plot_structure(valid,d_cm,title='')
        fig_bytes=None
        if fig:
            buf=BytesIO()
            fig.savefig(buf,format='png',dpi=150,bbox_inches='tight',facecolor='white')
            buf.seek(0); fig_bytes=buf.read()
            plt.close(fig)

        col_w=[3800,1400,3872]
        n_data=len(data_rows)
        tbl=doc.add_table(rows=1+n_data,cols=3)
        tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.LEFT

        hdr=tbl.rows[0]
        for i,cell in enumerate(hdr.cells): _set_cw(cell,col_w[i])
        _sc(hdr.cells[0],'รายละเอียด',bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,bg=HBG)
        _sc(hdr.cells[1],'หนา\n(ซม.)',bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,bg=HBG)
        _sc(hdr.cells[2],'ชนิดวัสดุ',bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,bg=HBG)

        for i,dr in enumerate(data_rows):
            row=tbl.rows[1+i]
            for j,cell in enumerate(row.cells): _set_cw(cell,col_w[j])
            lc=row.cells[0]
            if i==0:
                _set_vmerge(lc,restart=True); _cell_margin(lc)
                lc.text=''; p_img=lc.paragraphs[0]
                p_img.alignment=WD_ALIGN_PARAGRAPH.CENTER
                if fig_bytes:
                    p_img.add_run().add_picture(BytesIO(fig_bytes),width=Inches(2.4))
            else:
                _set_vmerge(lc,restart=False); lc.text=''
            _set_valign(lc,'center')
            _sc(row.cells[1],dr['thick'],align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_valign(row.cells[1],'center')
            _sc(row.cells[2],dr['material'])
            _set_valign(row.cells[2],'center')

        if fig_caption_text:
            _add_fig_caption(fig_caption_text)
        doc.add_paragraph()

    # ── JPCP section ──────────────────────────────────────────
    if inc_jpcp:
        _add_heading(f'{_sec_num(sec_prefix,1)}  ชั้นโครงสร้างทางคอนกรีตประเภท JPCP/JRCP',level=2)
        fig_n=nfig()
        _layer_table(jpcp_layers, jpcp_params.get('sel_d',30), 'JPCP/JRCP', jpcp_params.get('cbr',4.0))
        _esb_section(jpcp_layers)
        _add_heading(f'{_sec_num(sec_prefix,2)}  การคำนวณ Corrected Modulus of Subgrade Reaction (k-value) สำหรับ JPCP/JRCP',level=2)
        fig_n_k=nfig()
        _kvalue_section(jpcp_params, jpcp_fig33, jpcp_fig34, fig_n_k)
        if jpcp_fig34 and jpcp_params.get('ls',0)>0: fig_counter[0]+=1
        _add_heading('ผลการออกแบบความหนาผิวทางคอนกรีต JPCP/JRCP',level=3)
        _design_result(jpcp_params, jpcp_rows, jpcp_params.get('sel_d',30))

    # ── CRCP section ──────────────────────────────────────────
    if inc_crcp:
        sub_off=2 if inc_jpcp else 0
        _add_heading(f'{_sec_num(sec_prefix,sub_off+1)}  ชั้นโครงสร้างทางคอนกรีตประเภท CRCP',level=2)
        _layer_table(crcp_layers, crcp_params.get('sel_d',30), 'CRCP', crcp_params.get('cbr',4.0))
        _esb_section(crcp_layers)
        _add_heading(f'{_sec_num(sec_prefix,sub_off+2)}  การคำนวณ Corrected Modulus of Subgrade Reaction (k-value) สำหรับ CRCP',level=2)
        fig_n_k2=nfig()
        _kvalue_section(crcp_params, crcp_fig33, crcp_fig34, fig_n_k2)
        if crcp_fig34 and crcp_params.get('ls',0)>0: fig_counter[0]+=1
        _add_heading('ผลการออกแบบความหนาผิวทางคอนกรีต CRCP',level=3)
        _design_result(crcp_params, crcp_rows, crcp_params.get('sel_d',30))

    # ── Summary section ───────────────────────────────────────
    if inc_summary and (inc_jpcp or inc_crcp):
        doc.add_page_break()
        parts=sec_prefix.split('.')
        try: parts[-1]=str(int(parts[-1])+1); h_sum='.'.join(parts)
        except: h_sum=sec_prefix+'_สรุป'
        _add_heading(f'{h_sum}  สรุปโครงสร้างชั้นทางที่ออกแบบด้วยวิธี AASHTO 1993',level=1)
        _add_para(summary_text)
        doc.add_paragraph()
        pat=1
        if inc_jpcp:
            fn=nfig()
            _add_para(f'รูปแบบที่ {pat}: ผิวทางคอนกรีต แบบ JPCP/JRCP  (รูปที่ {fig_prefix}{fn})',bold=True)
            _summary_table(jpcp_layers, jpcp_params.get('sel_d',30), 'JPCP/JRCP',
                          jpcp_params.get('cbr',4.0),
                          f'รูปที่ {fig_prefix}{fn}  โครงสร้างชั้นทางรูปแบบที่ {pat} ผิวทางคอนกรีต แบบ JPCP/JRCP')
            pat+=1
        if inc_crcp:
            fn=nfig()
            _add_para(f'รูปแบบที่ {pat}: ผิวทางคอนกรีต แบบ CRCP  (รูปที่ {fig_prefix}{fn})',bold=True)
            _summary_table(crcp_layers, crcp_params.get('sel_d',30), 'CRCP',
                          crcp_params.get('cbr',4.0),
                          f'รูปที่ {fig_prefix}{fn}  โครงสร้างชั้นทางรูปแบบที่ {pat} ผิวทางคอนกรีต แบบ CRCP')

    # ── อ้างอิง ──────────────────────────────────────────────
    doc.add_paragraph()
    _add_para('เอกสารอ้างอิง',bold=True)
    _add_para('AASHTO Guide for Design of Pavement Structures 1993. American Association of State Highway and Transportation Officials, Washington, D.C.')

    buf=BytesIO(); doc.save(buf); buf.seek(0)
    return buf


# ============================================================
# Helpers — ดึงข้อมูลจาก session_state
# ============================================================
def _get_params(prefix):
    p = st.session_state.get(f'{prefix}_design_params', {})
    if not p: return None
    rec_d = st.session_state.get(f'{prefix}_rec_d_cm') or 30
    p['sel_d'] = rec_d
    p['cbr']   = st.session_state.get('cbr', 4.0)
    p['k_inf'] = st.session_state.get(f'{prefix}_k_inf', p.get('k_eff',0))
    p['DSB_in']= st.session_state.get(f'{prefix}_dsb', 0)
    p['ESB_psi']= st.session_state.get(f'{prefix}_esb', 0)
    p['MR_psi'] = st.session_state.get('MR_psi', 6000)
    p['ls']    = st.session_state.get(f'{prefix}_ls_val', 0.0)
    return p

def _has_data(prefix):
    return (st.session_state.get(f'{prefix}_design_params') is not None and
            st.session_state.get(f'{prefix}_design_rows') is not None and
            st.session_state.get(f'{prefix}_k_eff') is not None)


# ============================================================
# MAIN
# ============================================================
def render_tab4():
    has_j = _has_data('jpcp')
    has_c = _has_data('crcp')

    # ── Status ───────────────────────────────────────────────
    st.markdown('<div class="rp-card">', unsafe_allow_html=True)
    st.markdown('<div class="rp-card-title">📋 สถานะข้อมูล</div>', unsafe_allow_html=True)
    c1,c2,c3 = st.columns(3)
    with c1:
        if has_j: st.markdown('<div class="rp-status-ok">✅ JPCP/JRCP — พร้อม export</div>', unsafe_allow_html=True)
        else:     st.markdown('<div class="rp-status-warn">⚠️ JPCP/JRCP — ยังไม่ได้คำนวณ (Tab 3)</div>', unsafe_allow_html=True)
    with c2:
        if has_c: st.markdown('<div class="rp-status-ok">✅ CRCP — พร้อม export</div>', unsafe_allow_html=True)
        else:     st.markdown('<div class="rp-status-warn">⚠️ CRCP — ยังไม่ได้คำนวณ (Tab 3)</div>', unsafe_allow_html=True)
    with c3:
        proj = st.session_state.get('project_name','') or '—'
        st.markdown(f'<div class="rp-status-info">📁 โครงการ: {proj}</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    if not has_j and not has_c:
        st.warning('⚠️ กรุณาคำนวณใน Tab 3 ก่อนอย่างน้อย 1 ประเภท')
        return

    # ── Layout: ซ้าย=ตั้งค่า / ขวา=preview ──────────────────
    col_cfg, col_prev = st.columns([1.1, 0.9])

    with col_cfg:
        st.markdown('### ⚙️ ตั้งค่ารายงาน')

        # เลขหัวข้อ
        with st.expander('🔢 เลขหัวข้อและเลขรูป', expanded=True):
            c1,c2,c3 = st.columns([1,1,0.8])
            with c1:
                sec_prefix = st.text_input('Prefix หัวข้อหลัก (เช่น 4.5)',
                                           value=st.session_state.get('rpt_sec_prefix','4.5'),
                                           key='rpt_sec_prefix')
            with c2:
                fig_prefix = st.text_input('Prefix เลขรูป (เช่น 4-)',
                                           value=st.session_state.get('rpt_fig_prefix','4-'),
                                           key='rpt_fig_prefix')
            with c3:
                fig_start = st.number_input('เริ่มที่รูปที่',
                                            min_value=1, max_value=99,
                                            value=st.session_state.get('rpt_fig_start',5),
                                            key='rpt_fig_start')
            st.caption(f'ตัวอย่าง: รูปที่ {fig_prefix}{fig_start}, {fig_prefix}{fig_start+1} ...')

        # บทเกริ่นนำ
        with st.expander('📄 บทเกริ่นนำ', expanded=True):
            intro = st.text_area('เนื้อหาบทเกริ่นนำ (แก้ไขได้)',
                                 value=st.session_state.get('rpt_intro', DEFAULT_INTRO),
                                 height=160, key='rpt_intro')

        # บทสรุป
        with st.expander('📋 บทสรุป (หัวข้อสรุปโครงสร้างชั้นทาง)', expanded=False):
            # generate อัตโนมัติ
            auto_summary = DEFAULT_SUMMARY
            if has_j:
                jp = _get_params('jpcp')
                auto_summary += f'\n\nรูปแบบที่ 1 ผิวทางคอนกรีต แบบ JPCP/JRCP ความหนา {jp.get("sel_d",30)} ซม. ({round(jp.get("sel_d",30)/2.54)} นิ้ว) k_eff = {jp.get("k_eff",0):.0f} pci'
            if has_c:
                cp = _get_params('crcp')
                pat = 2 if has_j else 1
                auto_summary += f'\n\nรูปแบบที่ {pat} ผิวทางคอนกรีต แบบ CRCP ความหนา {cp.get("sel_d",30)} ซม. ({round(cp.get("sel_d",30)/2.54)} นิ้ว) k_eff = {cp.get("k_eff",0):.0f} pci'

            summary = st.text_area('เนื้อหาบทสรุป (แก้ไขได้)',
                                   value=st.session_state.get('rpt_summary', auto_summary),
                                   height=140, key='rpt_summary')
            inc_summary = st.checkbox('รวมหัวข้อสรุปในรายงาน', value=True, key='rpt_inc_sum')

        # เลือก export
        st.markdown('---')
        st.markdown('### 📥 Export รายงาน')
        calc_date = datetime.now().strftime('%d/%m/%Y')
        proj_name = st.session_state.get('project_name','')

        def _make_report(inc_j, inc_c, label_suffix=''):
            jp_p = _get_params('jpcp') if inc_j else {}
            cp_p = _get_params('crcp') if inc_c else {}
            try:
                buf = _build_report(
                    sec_prefix, fig_prefix, int(fig_start),
                    intro, DEFAULT_PAVEMENT_DESC, summary,
                    proj_name, calc_date,
                    inc_j,
                    st.session_state.get('jpcp_layers',[]) if inc_j else [],
                    jp_p, st.session_state.get('jpcp_design_rows',[]) if inc_j else [],
                    st.session_state.get('jpcp_fig33_bytes') if inc_j else None,
                    st.session_state.get('jpcp_fig34_bytes') if inc_j else None,
                    inc_c,
                    st.session_state.get('crcp_layers',[]) if inc_c else [],
                    cp_p, st.session_state.get('crcp_design_rows',[]) if inc_c else [],
                    st.session_state.get('crcp_fig33_bytes') if inc_c else None,
                    st.session_state.get('crcp_fig34_bytes') if inc_c else None,
                    inc_summary,
                )
            except Exception as e:
                st.error(f'❌ {e}')
                return
            if buf is None:
                st.error('❌ ไม่พบ python-docx')
                return
            fname = f'Report_{label_suffix}_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'
            st.download_button(
                f'📥 ดาวน์โหลด {label_suffix} (.docx)', buf, fname,
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                key=f'dl_{label_suffix}', use_container_width=True)

        if has_j and has_c:
            _make_report(True, True,  'Combined_JPCP_CRCP')
            st.markdown('<div style="height:6px"></div>', unsafe_allow_html=True)
            cc1, cc2 = st.columns(2)
            with cc1: _make_report(True,  False, 'JPCP_only')
            with cc2: _make_report(False, True,  'CRCP_only')
        elif has_j:
            _make_report(True, False, 'JPCP_only')
        elif has_c:
            _make_report(False, True, 'CRCP_only')

    # ── Preview ───────────────────────────────────────────────
    with col_prev:
        st.markdown('### 👁️ ตัวอย่างโครงสร้างรายงาน')
        fn = int(fig_start)
        sub_j1 = f'{sec_prefix}.1'; sub_j2 = f'{sec_prefix}.2'
        sub_c_off = 2 if has_j else 0
        sub_c1 = f'{sec_prefix}.{sub_c_off+1}'; sub_c2 = f'{sec_prefix}.{sub_c_off+2}'
        parts_sec = sec_prefix.split('.')
        try: parts_sec[-1]=str(int(parts_sec[-1])+1); sum_sec='.'.join(parts_sec)
        except: sum_sec=sec_prefix+'_สรุป'

        preview_lines = [('page','หน้าปก')]
        preview_lines.append(('h1', f'{sec_prefix} การออกแบบผิวทางคอนกรีต'))
        preview_lines.append(('note','บทเกริ่นนำ + สมการ AASHTO 1993'))

        if has_j:
            jd = _get_params('jpcp')
            jsel = jd.get('sel_d',30) if jd else 30
            preview_lines.append(('h2', f'{sub_j1} ชั้นโครงสร้างทาง JPCP/JRCP'))
            preview_lines.append(('note', f'รูปที่ {fig_prefix}{fn} + ตาราง ESB'))
            fn+=1
            preview_lines.append(('h2', f'{sub_j2} k-value สำหรับ JPCP/JRCP'))
            preview_lines.append(('note', f'Nomograph + ตาราง k_eff + ผลการออกแบบ (D={jsel} ซม.)'))
            ls_j = st.session_state.get('jpcp_ls_val',0)
            if ls_j>0: fn+=1
            fn+=1

        if has_c:
            cd2 = _get_params('crcp')
            csel = cd2.get('sel_d',30) if cd2 else 30
            preview_lines.append(('h2', f'{sub_c1} ชั้นโครงสร้างทาง CRCP'))
            preview_lines.append(('note', f'รูปที่ {fig_prefix}{fn} + ตาราง ESB'))
            fn+=1
            preview_lines.append(('h2', f'{sub_c2} k-value สำหรับ CRCP'))
            preview_lines.append(('note', f'Nomograph + ตาราง k_eff + ผลการออกแบบ (D={csel} ซม.)'))
            ls_c = st.session_state.get('crcp_ls_val',0)
            if ls_c>0: fn+=1
            fn+=1

        if inc_summary and (has_j or has_c):
            preview_lines.append(('h1', f'{sum_sec} สรุปโครงสร้างชั้นทาง AASHTO 1993'))
            pat=1
            if has_j:
                preview_lines.append(('note', f'รูปแบบที่ {pat}: JPCP/JRCP (รูปที่ {fig_prefix}{fn})'))
                fn+=1; pat+=1
            if has_c:
                preview_lines.append(('note', f'รูปแบบที่ {pat}: CRCP (รูปที่ {fig_prefix}{fn})'))

        # render preview
        html = '<div style="background:#FAFAFA;border:1.5px solid #E0E0E0;border-radius:10px;padding:14px 16px;font-size:13px;line-height:1.8">'
        for kind, text in preview_lines:
            if kind == 'page':
                html += f'<div style="border-bottom:1px solid #BDBDBD;padding-bottom:6px;margin-bottom:8px;font-weight:700;color:#37474F">📄 {text}</div>'
            elif kind == 'h1':
                html += f'<div style="font-weight:700;color:#1565C0;margin-top:8px;text-decoration:underline">{text}</div>'
            elif kind == 'h2':
                html += f'<div style="font-weight:600;color:#37474F;margin-left:14px">{text}</div>'
            elif kind == 'note':
                html += f'<div style="color:#78909C;margin-left:28px;font-size:12px;font-style:italic">{text}</div>'
        html += '</div>'
        st.markdown(html, unsafe_allow_html=True)

        st.markdown('<div style="height:8px"></div>', unsafe_allow_html=True)
        st.caption('🔴 หมายเหตุ: รายงานใช้ข้อมูลจาก Tab 1 (ESAL) Tab 2 (Subgrade) และ Tab 3 (Design)')
